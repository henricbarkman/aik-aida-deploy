"""AIda Web UI. Flask app with split chat/results layout."""

from __future__ import annotations

import json
import os
import secrets
import sys
from functools import wraps
from urllib.error import HTTPError
from urllib.parse import urlencode
from urllib.request import Request, urlopen

from flask import Flask, Response, jsonify, redirect, render_template_string, request, session, url_for

# Add src to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', '..'))

import anthropic

from aida.agents.aggregate import compute_aggregate
from aida.agents.alternatives import find_alternatives
from aida.agents.baseline import calculate_baseline
from aida.agents.intake import run_intake
from aida.agents.report import generate_report_markdown
from aida.models import Baseline, Project, Selections

# Timeout errors to catch from Anthropic SDK
_TIMEOUT_ERRORS = (anthropic.APITimeoutError, TimeoutError)

app = Flask(__name__)
app.secret_key = os.environ.get('AIDA_SECRET_KEY', secrets.token_hex(32))

AIDA_PASSWORD = os.environ.get('AIDA_PASSWORD', '')
SUPABASE_URL = os.environ.get('SUPABASE_URL', '').strip()
SUPABASE_ANON_KEY = os.environ.get('SUPABASE_ANON_KEY', '').strip()
SUPABASE_JWT_SECRET = os.environ.get('SUPABASE_JWT_SECRET', '').strip()

try:
    import jwt as pyjwt
    from jwt import PyJWKClient
except ImportError:
    pyjwt = None
    PyJWKClient = None

# JWKS client for ES256 token verification (cached, Supabase default since 2026)
_jwks_client = None


def _get_jwks_client():
    global _jwks_client
    if _jwks_client is None and PyJWKClient and SUPABASE_URL:
        _jwks_client = PyJWKClient(f"{SUPABASE_URL}/auth/v1/.well-known/jwks.json")
    return _jwks_client


def get_user_from_token():
    """Extract user_id from Supabase JWT in Authorization header."""
    if not pyjwt or not SUPABASE_URL:
        return None
    auth_header = request.headers.get('Authorization', '')
    if not auth_header.startswith('Bearer '):
        return None
    token = auth_header[7:]

    # Try ES256 via JWKS first (Supabase default since 2024)
    jwks = _get_jwks_client()
    if jwks:
        try:
            signing_key = jwks.get_signing_key_from_jwt(token)
            payload = pyjwt.decode(
                token, signing_key.key,
                algorithms=['ES256'], audience='authenticated'
            )
            return payload.get('sub')
        except Exception as e:
            app.logger.debug("ES256 JWKS validation failed: %s", e)

    # Fallback: HS256 with local secret
    if SUPABASE_JWT_SECRET:
        try:
            payload = pyjwt.decode(
                token, SUPABASE_JWT_SECRET,
                algorithms=['HS256'], audience='authenticated'
            )
            return payload.get('sub')
        except Exception as e:
            app.logger.debug("HS256 validation failed: %s", e)

    # Last resort: verify token via Supabase auth API (handles any algorithm)
    try:
        resp = __import__('urllib.request', fromlist=['urlopen']).urlopen(
            __import__('urllib.request', fromlist=['Request']).Request(
                f"{SUPABASE_URL}/auth/v1/user",
                headers={
                    'apikey': SUPABASE_ANON_KEY,
                    'Authorization': f'Bearer {token}',
                },
            ),
            timeout=5,
        )
        user_data = json.loads(resp.read().decode())
        uid = user_data.get('id')
        if uid:
            app.logger.info("Token validated via Supabase /auth/v1/user fallback")
            return uid
    except Exception as e:
        app.logger.debug("Supabase /auth/v1/user fallback failed: %s", e)

    return None


def supabase_request(method, path, data=None, token=None, params=None):
    """Make a request to Supabase REST API (PostgREST)."""
    url = f"{SUPABASE_URL}/rest/v1/{path}"
    if params:
        url += '?' + urlencode(params)
    headers = {
        'apikey': SUPABASE_ANON_KEY,
        'Content-Type': 'application/json',
        'Prefer': 'return=representation',
    }
    if token:
        headers['Authorization'] = f'Bearer {token}'
    body = json.dumps(data).encode() if data else None
    req = Request(url, data=body, headers=headers, method=method)
    try:
        with urlopen(req) as resp:
            resp_data = resp.read().decode()
            return json.loads(resp_data) if resp_data else None
    except HTTPError as e:
        error_body = e.read().decode()
        raise Exception(f"Supabase error {e.code}: {error_body}")


def require_auth(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        # Supabase JWT auth
        if SUPABASE_URL:
            user_id = get_user_from_token()
            if not user_id:
                return jsonify({'error': 'Ej inloggad'}), 401
            request.user_id = user_id
            return f(*args, **kwargs)
        # Legacy password auth
        if not AIDA_PASSWORD:
            return f(*args, **kwargs)
        if session.get('authenticated'):
            return f(*args, **kwargs)
        if request.is_json:
            return jsonify({'error': 'Ej inloggad'}), 401
        return redirect(url_for('login'))
    return decorated


def require_supabase_auth(f):
    """Like require_auth but only allows Supabase JWT (for CRUD endpoints)."""
    @wraps(f)
    def decorated(*args, **kwargs):
        if not SUPABASE_URL:
            return jsonify({'error': 'Supabase ej konfigurerat'}), 501
        user_id = get_user_from_token()
        if not user_id:
            return jsonify({'error': 'Ej inloggad'}), 401
        request.user_id = user_id
        return f(*args, **kwargs)
    return decorated


LOGIN_TEMPLATE = r"""<!DOCTYPE html>
<html lang="sv">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>AIda | Logga in</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=Roboto:wght@400;500;700&display=swap" rel="stylesheet">
<style>
:root { --kk-gold: #FFCC01; --kk-dark-red: #B5201F; --kk-burgundy: #890200; --kk-charcoal: #444; --kk-cream: #FFF9DE; --kk-warm-bg: #FFFBF5; --kk-gray-200: #e5e5e5; --kk-gray-400: #a3a3a3; --kk-gold-light: #FFF1B6; }
* { margin: 0; padding: 0; box-sizing: border-box; }
body { font-family: 'Roboto', sans-serif; height: 100vh; display: flex; align-items: center; justify-content: center; background: var(--kk-warm-bg); }
.login-box { background: white; border-radius: 12px; padding: 40px; width: 360px; box-shadow: 0 4px 24px rgba(0,0,0,0.08); border-top: 3px solid var(--kk-gold-light); }
.login-box h1 { font-size: 24px; color: var(--kk-charcoal); margin-bottom: 8px; }
.login-box p { font-size: 13px; color: var(--kk-gray-400); margin-bottom: 24px; }
.login-box input { width: 100%; padding: 12px 16px; border: 1px solid var(--kk-gray-200); border-radius: 8px; font-size: 14px; font-family: inherit; outline: none; }
.login-box input:focus { border-color: var(--kk-dark-red); box-shadow: 0 0 0 2px rgba(181,32,31,0.15); }
.login-box button { width: 100%; padding: 12px; background: var(--kk-charcoal); color: white; border: none; border-radius: 8px; font-size: 14px; font-weight: 600; cursor: pointer; margin-top: 12px; font-family: inherit; }
.login-box button:hover { background: var(--kk-dark-red); }
.error { color: var(--kk-dark-red); font-size: 12px; margin-top: 8px; }
.footer { position: fixed; bottom: 16px; font-size: 11px; color: var(--kk-gray-400); }
</style>
</head>
<body>
<div class="login-box">
  <h1>AIda</h1>
  <p>Klimatkalkyl och beslutsstöd för ombyggnationer</p>
  <form method="POST">
    <input type="password" name="password" placeholder="Lösenord" autofocus>
    {% if error %}<div class="error">{{ error }}</div>{% endif %}
    <button type="submit">Logga in</button>
  </form>
</div>
<div class="footer"></div>
</body>
</html>"""


@app.route('/login', methods=['GET', 'POST'])
def login():
    if not AIDA_PASSWORD:
        return redirect(url_for('index'))
    error = None
    if request.method == 'POST':
        if request.form.get('password') == AIDA_PASSWORD:
            session['authenticated'] = True
            return redirect(url_for('index'))
        error = 'Fel lösenord'
    return render_template_string(LOGIN_TEMPLATE, error=error)


@app.route('/')
def index():
    if SUPABASE_URL:
        return render_template_string(HTML_TEMPLATE,
            supabase_url=SUPABASE_URL,
            supabase_anon_key=SUPABASE_ANON_KEY,
            has_supabase=True)
    if AIDA_PASSWORD and not session.get('authenticated'):
        return redirect(url_for('login'))
    return render_template_string(HTML_TEMPLATE,
        supabase_url='', supabase_anon_key='', has_supabase=False)


@app.route('/docs/<path:filename>')
def serve_docs(filename):
    """Serve static docs files."""
    # Resolve relative to this file: src/aida/web/app.py -> project_root/docs/
    docs_dir = os.path.abspath(os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..', '..', 'docs'))
    filepath = os.path.abspath(os.path.join(docs_dir, filename))
    if not filepath.startswith(docs_dir):
        return 'Forbidden', 403
    try:
        with open(filepath) as f:
            return f.read(), 200, {'Content-Type': 'text/html; charset=utf-8'}
    except FileNotFoundError:
        return 'Not found', 404


@app.route('/api/intake', methods=['POST'])
@require_auth
def api_intake():
    data = request.json
    description = data.get('description', '')
    if not description:
        return jsonify({'error': 'Beskrivning saknas'}), 400

    try:
        result = run_intake(description)
        return jsonify(result)
    except _TIMEOUT_ERRORS:
        return jsonify({'error': 'Analysen tog för lång tid. Försök igen.'}), 504
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/baseline', methods=['POST'])
@require_auth
def api_baseline():
    data = request.json
    project_data = data.get('project')
    component_ids = data.get('component_ids')

    if not project_data:
        return jsonify({'error': 'Projekt saknas'}), 400

    try:
        project = Project.from_dict(project_data)

        # Partial rerun: filter the project to a subset of components before LLM matching.
        # An empty/missing component_ids means "all components" (full rerun, original behavior).
        if isinstance(component_ids, list) and component_ids:
            requested = set(component_ids)
            project.components = [c for c in project.components if c.id in requested]
            if not project.components:
                return jsonify({'error': 'Inga komponenter matchade angivna component_ids'}), 400

        baseline = calculate_baseline(project)
        return jsonify(baseline.to_dict())
    except _TIMEOUT_ERRORS:
        return jsonify({'error': 'Baslinjeberäkningen tog för lång tid. Försök igen.'}), 504
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/alternatives', methods=['POST'])
@require_auth
def api_alternatives():
    data = request.json
    project_data = data.get('project')
    baseline_data = data.get('baseline')
    component_ids = data.get('component_ids')

    if not project_data or not baseline_data:
        return jsonify({'error': 'Projekt eller baslinje saknas'}), 400

    try:
        project = Project.from_dict(project_data)
        baseline = Baseline.from_dict(baseline_data)

        # Partial rerun: filter both project and baseline to the requested subset.
        # find_alternatives iterates over baseline.components, so both must agree on which
        # components are in scope. Empty/missing list = all components (full rerun).
        if isinstance(component_ids, list) and component_ids:
            requested = set(component_ids)
            project_ids_before = {c.id for c in project.components}
            baseline_ids_before = {c.component_id for c in baseline.components}
            project.components = [c for c in project.components if c.id in requested]
            baseline.components = [c for c in baseline.components if c.component_id in requested]
            if not project.components:
                missing = sorted(requested - project_ids_before)
                return jsonify({'error': f'Komponent saknas i projektet: {missing}'}), 400
            if not baseline.components:
                missing = sorted(requested - baseline_ids_before)
                return jsonify({
                    'error': f'Komponent saknas i baslinjen: {missing}. Kör om baslinjen först.'
                }), 400

        user_feedback = data.get('user_feedback')
        result = find_alternatives(project, baseline, user_feedback=user_feedback)
        return jsonify(result.to_dict())
    except _TIMEOUT_ERRORS:
        return jsonify({'error': 'Alternativanalysen tog för lång tid. Försök igen, eller minska antal komponenter.'}), 504
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/aggregate', methods=['POST'])
@require_auth
def api_aggregate():
    data = request.json
    try:
        project = Project.from_dict(data.get('project', {}))
        selections = Selections.from_dict(data.get('selections', {}))
        result = compute_aggregate(project, selections)
        return jsonify(result.to_dict())
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/report', methods=['POST'])
@require_auth
def api_report():
    data = request.json
    try:
        project = Project.from_dict(data.get('project', {}))
        selections = Selections.from_dict(data.get('selections', {}))
        if not selections.components:
            return jsonify({'error': 'Inga komponenter valda'}), 400
        markdown = generate_report_markdown(project, selections)
        return jsonify({'markdown': markdown})
    except _TIMEOUT_ERRORS:
        return jsonify({'error': 'Rapportgenereringen tog för lång tid. Försök igen.'}), 504
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/report/docx', methods=['POST'])
@require_auth
def api_report_docx():
    """Convert markdown report to .docx and return as download."""
    data = request.json or {}
    markdown = data.get('markdown', '')
    if not markdown:
        return jsonify({'error': 'Markdown saknas'}), 400

    try:
        import io
        import re
        from datetime import date

        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt, RGBColor

        BRAND_BLUE = RGBColor(0x4A, 0x90, 0xD9)
        GRAY_66 = RGBColor(0x66, 0x66, 0x66)
        GRAY_99 = RGBColor(0x99, 0x99, 0x99)
        HEADER_BG = "4A90D9"

        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # Base style
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)

        # Heading styles
        for level in range(1, 4):
            h_style = doc.styles[f'Heading {level}']
            h_style.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

        # AIda branding header
        header_p = doc.add_paragraph()
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header_p.add_run('AIda')
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = BRAND_BLUE
        run = header_p.add_run('  |  Klimatberäkning av ombyggnad')
        run.font.size = Pt(10)
        run.font.color.rgb = GRAY_66

        active_table = None
        table_is_first_row = False

        def _add_rich_runs(paragraph, text):
            """Parse inline markdown (bold, italic) into runs on a paragraph."""
            # Split on bold (**text**) and italic (*text*) markers
            parts = re.split(r'(\*\*[^*]+?\*\*|\*[^*]+?\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = paragraph.add_run(part[2:-2])
                    r.bold = True
                elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                    r = paragraph.add_run(part[1:-1])
                    r.italic = True
                else:
                    paragraph.add_run(part)

        def _add_rich_paragraph(text, style_name=None):
            """Add paragraph with bold/italic markdown spans preserved."""
            p = doc.add_paragraph(style=style_name)
            _add_rich_runs(p, text)
            return p

        def _style_header_cell(cell):
            """Apply white-on-blue header styling to a table cell."""
            from docx.oxml.ns import qn
            shading = cell._element.find(qn('w:tcPr'))
            if shading is None:
                tc_pr = cell._element.makeelement(qn('w:tcPr'), {})
                cell._element.insert(0, tc_pr)
            else:
                tc_pr = shading
            shading_el = tc_pr.makeelement(qn('w:shd'), {
                qn('w:val'): 'clear',
                qn('w:color'): 'auto',
                qn('w:fill'): HEADER_BG,
            })
            tc_pr.append(shading_el)
            for paragraph in cell.paragraphs:
                for r in paragraph.runs:
                    r.bold = True
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    r.font.size = Pt(10)

        for line in markdown.split('\n'):
            stripped = line.strip()

            if stripped.startswith('#### '):
                active_table = None
                p = doc.add_paragraph()
                run = p.add_run(stripped[5:])
                run.bold = True
                run.font.size = Pt(11)
            elif stripped.startswith('### '):
                active_table = None
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith('## '):
                active_table = None
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith('# '):
                active_table = None
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith('- ') or stripped.startswith('* '):
                active_table = None
                _add_rich_paragraph(stripped[2:], style_name='List Bullet')
            elif re.match(r'^\d+\.\s', stripped):
                active_table = None
                text = re.sub(r'^\d+\.\s', '', stripped)
                _add_rich_paragraph(text, style_name='List Number')
            elif stripped.startswith('|') and '|' in stripped[1:]:
                cells = [c.strip() for c in stripped.split('|')[1:-1]]
                # Skip separator rows (|---|---|)
                if cells and all(set(c) <= {'-', ':', ' '} for c in cells):
                    continue
                if cells:
                    if active_table is None:
                        active_table = doc.add_table(rows=0, cols=len(cells))
                        active_table.style = 'Table Grid'
                        active_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        table_is_first_row = True
                    if len(cells) == len(active_table.columns):
                        row = active_table.add_row()
                        for i, cell_text in enumerate(cells):
                            cell_text = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
                            row.cells[i].text = cell_text
                            # Smaller font in table cells
                            for paragraph in row.cells[i].paragraphs:
                                for r in paragraph.runs:
                                    r.font.size = Pt(10)
                            if table_is_first_row:
                                _style_header_cell(row.cells[i])
                        table_is_first_row = False
            elif stripped == '':
                active_table = None
            elif stripped.startswith('---') or stripped.startswith('***'):
                # Horizontal rule -- skip
                active_table = None
            else:
                active_table = None
                _add_rich_paragraph(stripped)

        # Footer with disclaimer
        doc.add_paragraph()
        footer_p = doc.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.add_run('Genererad av AIda | AI-stödd klimatanalys för ombyggnadsprojekt')
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY_99

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        today = date.today().strftime('%Y-%m-%d')
        filename = f'AIda_rapport_{today}.docx'

        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f'attachment; filename={filename}'},
        )
    except ImportError:
        return jsonify({'error': 'python-docx är inte installerat på servern'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/chat', methods=['POST'])
@require_auth
def api_chat():
    """Conversational endpoint with tool-use.

    Request:
      {
        message: str,
        history: [{role, content}],
        project?: Project,
        baseline?: Baseline,
        alternatives?: AlternativesResult,
        selections?: {component_id: Selection}
      }

    Response:
      {
        reply: str,
        state_updates?: {project?, baseline?, alternatives?, selections?},
        tool_calls?: [...]
      }
    """
    from aida.agents.chat_agent import run_chat_agent

    data = request.json or {}
    try:
        result = run_chat_agent(
            message=data.get('message', ''),
            history=data.get('history', []),
            project=data.get('project'),
            baseline=data.get('baseline'),
            alternatives=data.get('alternatives'),
            selections=data.get('selections'),
        )
        return jsonify(result)
    except _TIMEOUT_ERRORS:
        return jsonify({'error': 'Chatten svarade inte i tid. Försök igen.'}), 504
    except Exception as e:
        app.logger.exception("chat_agent failed")
        return jsonify({'error': str(e)}), 500


# === Analyses CRUD (Supabase) ===

@app.route('/api/analyses', methods=['POST'])
@require_supabase_auth
def create_analysis():
    token = request.headers.get('Authorization', '').replace('Bearer ', '')
    data = request.json or {}
    row = {
        'user_id': request.user_id,
        'name': data.get('name', 'Nytt projekt'),
        'status': data.get('status', 'intake'),
        'project_data': data.get('project_data'),
        'baseline_data': data.get('baseline_data'),
        'alternatives_data': data.get('alternatives_data'),
        'selections_data': data.get('selections_data'),
        'report_markdown': data.get('report_markdown'),
    }
    result = supabase_request('POST', 'analyses', data=row, token=token)
    return jsonify(result[0] if isinstance(result, list) else result)


@app.route('/api/analyses', methods=['GET'])
@require_supabase_auth
def list_analyses():
    token = request.headers.get('Authorization', '').replace('Bearer ', '')
    params = {
        'select': 'id,name,status,created_at,updated_at',
        'user_id': f'eq.{request.user_id}',
        'order': 'updated_at.desc',
        'limit': '20',
    }
    result = supabase_request('GET', 'analyses', token=token, params=params)
    return jsonify(result or [])


@app.route('/api/analyses/<analysis_id>', methods=['GET'])
@require_supabase_auth
def get_analysis(analysis_id):
    token = request.headers.get('Authorization', '').replace('Bearer ', '')
    params = {
        'id': f'eq.{analysis_id}',
        'user_id': f'eq.{request.user_id}',
    }
    result = supabase_request('GET', 'analyses', token=token, params=params)
    if not result:
        return jsonify({'error': 'Ej hittad'}), 404
    return jsonify(result[0])


@app.route('/api/analyses/<analysis_id>', methods=['PUT'])
@require_supabase_auth
def update_analysis(analysis_id):
    token = request.headers.get('Authorization', '').replace('Bearer ', '')
    data = request.json or {}
    update = {}
    for key in ('name', 'status', 'project_data', 'baseline_data',
                'alternatives_data', 'selections_data', 'report_markdown'):
        if key in data:
            update[key] = data[key]
    params = {
        'id': f'eq.{analysis_id}',
        'user_id': f'eq.{request.user_id}',
    }
    result = supabase_request('PATCH', 'analyses', data=update, token=token, params=params)
    if not result:
        return jsonify({'error': 'Ej hittad'}), 404
    return jsonify(result[0] if isinstance(result, list) else result)


@app.route('/api/analyses/<analysis_id>', methods=['DELETE'])
@require_supabase_auth
def delete_analysis(analysis_id):
    token = request.headers.get('Authorization', '').replace('Bearer ', '')
    params = {
        'id': f'eq.{analysis_id}',
        'user_id': f'eq.{request.user_id}',
    }
    result = supabase_request('DELETE', 'analyses', token=token, params=params)
    if not result:
        return jsonify({'error': 'Ej hittad'}), 404
    return jsonify({'ok': True})


HTML_TEMPLATE = r"""<!DOCTYPE html>
<html lang="sv">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>AIda | Klimatkalkyl för ombyggnationer</title>
<link rel="icon" href="data:image/svg+xml,<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='%23E84E0D' stroke-width='1.5' stroke-linecap='round'><circle cx='12' cy='12' r='5'/><path d='M12 1v3M12 20v3M1 12h3M20 12h3M4.2 4.2l2.1 2.1M17.7 17.7l2.1 2.1M4.2 19.8l2.1-2.1M17.7 6.3l2.1-2.1'/></svg>">
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Roboto:wght@400;500;700&display=swap" rel="stylesheet">
<script src="https://cdn.jsdelivr.net/npm/marked/marked.min.js"></script>
<script src="https://cdn.jsdelivr.net/npm/dompurify@3/dist/purify.min.js"></script>
{% if has_supabase %}<script src="https://cdn.jsdelivr.net/npm/@supabase/supabase-js@2"></script>{% endif %}
<style>
/* === Karlstads kommun färgpalett (karlstad.se-manér) === */
:root {
  --kk-gold: #FFCC01;
  --kk-gold-light: #FFF1B6;
  --kk-orange: #EF7D00;
  --kk-red-orange: #E84E0D;
  --kk-red: #D41318;
  --kk-dark-red: #B5201F;
  --kk-burgundy: #890200;
  --kk-cream: #FFF9DE;
  --kk-warm-bg: #FFFBF5;
  --kk-charcoal: #444444;
  --kk-text: #444444;
  --kk-gray-50: #fafafa;
  --kk-gray-100: #f5f5f5;
  --kk-gray-200: #e5e5e5;
  --kk-gray-300: #d4d4d4;
  --kk-gray-400: #a3a3a3;
  --kk-gray-500: #737373;
  --green-saving: #4a7c59;
}

* { margin: 0; padding: 0; box-sizing: border-box; }
body { font-family: 'Roboto', -apple-system, BlinkMacSystemFont, sans-serif; height: 100vh; display: flex; flex-direction: column; background: white; color: var(--kk-text); }

/* === Top bar (karlstad.se: white with warm accent line) === */
.topbar { background: white; color: var(--kk-charcoal); height: 56px; display: flex; align-items: center; justify-content: space-between; padding: 0 24px; flex-shrink: 0; border-bottom: 3px solid var(--kk-gold-light); }
.topbar-logo { display: flex; align-items: center; gap: 10px; }
.topbar-logo svg { width: 28px; height: 28px; color: var(--kk-red-orange); }
.topbar-logo span { font-size: 16px; font-weight: 700; letter-spacing: 0.5px; color: var(--kk-charcoal); }
.topbar-center { font-size: 14px; color: var(--kk-gray-500); }
.topbar-right { font-size: 12px; color: var(--kk-gray-400); }

/* === Progress tracker (mockup: numbered circles with line) === */
.progress-bar { padding: 24px 48px 16px; flex-shrink: 0; }
.progress-track { display: flex; justify-content: space-between; align-items: flex-start; position: relative; }
.progress-line { position: absolute; top: 16px; left: 40px; right: 40px; height: 2px; background: var(--kk-gray-200); }
.progress-fill { position: absolute; top: 0; left: 0; height: 100%; background: var(--kk-charcoal); transition: width 0.5s ease; }
.step-item { display: flex; flex-direction: column; align-items: center; z-index: 1; min-width: 80px; }
.step-circle { width: 32px; height: 32px; border-radius: 50%; display: flex; align-items: center; justify-content: center; font-size: 14px; font-weight: 600; transition: all 0.3s; background: white; color: var(--kk-gray-400); border: 2px solid var(--kk-gray-200); }
.step-circle.active { background: var(--kk-charcoal); color: white; border-color: var(--kk-charcoal); box-shadow: 0 2px 8px rgba(68,68,68,0.4); transform: scale(1.1); }
.step-circle.done { background: var(--kk-charcoal); color: white; border-color: var(--kk-charcoal); }
.step-label { margin-top: 6px; font-size: 11px; font-weight: 500; color: var(--kk-gray-500); text-align: center; }
.step-label.active { color: var(--kk-charcoal); font-weight: 700; }
.step-label.done { color: var(--kk-charcoal); }

/* === Main layout === */
.main { display: flex; flex: 1; overflow: hidden; padding: 0 24px 0 24px; gap: 24px; }

/* === Chat panel (mockup: rounded, warm bg) === */
.chat-panel { width: 40%; display: flex; flex-direction: column; flex-shrink: 0; }
.chat-container { flex: 1; display: flex; flex-direction: column; background: var(--kk-warm-bg); border-radius: 12px; border: 1px solid var(--kk-gray-200); overflow: hidden; min-height: 0; }
.chat-header { padding: 10px 16px; border-bottom: 1px solid var(--kk-gray-200); background: var(--kk-cream); display: flex; justify-content: space-between; align-items: center; }
.chat-header h2 { font-size: 15px; font-weight: 600; color: var(--kk-charcoal); }
.messages { flex: 1; overflow-y: auto; padding: 16px; display: flex; flex-direction: column; gap: 10px; }
.msg { padding: 10px 14px; border-radius: 16px; max-width: 85%; font-size: 13px; line-height: 1.5; }
.msg.user { background: #FFF0D4; color: var(--kk-text); align-self: flex-end; border-bottom-right-radius: 4px; }
.msg.bot { background: white; color: var(--kk-text); align-self: flex-start; border-bottom-left-radius: 4px; box-shadow: 0 1px 2px rgba(0,0,0,0.05); }
@keyframes msgIn { from { opacity: 0; transform: translateY(8px); } to { opacity: 1; transform: translateY(0); } }
.msg { animation: msgIn 0.25s ease-out; }
.msg.system { background: var(--kk-cream); font-size: 12px; text-align: center; align-self: center; max-width: 100%; color: var(--kk-gray-500); border: 1px solid var(--kk-gray-200); }
.msg p { margin: 0 0 8px; }
.msg p:last-child { margin-bottom: 0; }
.msg ol, .msg ul { margin: 6px 0; padding-left: 20px; }
.msg li { margin-bottom: 4px; }
.msg h1, .msg h2, .msg h3, .msg h4 { margin: 8px 0 4px; line-height: 1.3; }
.msg h1 { font-size: 16px; } .msg h2 { font-size: 15px; } .msg h3 { font-size: 14px; } .msg h4 { font-size: 13px; }
.msg code { background: rgba(0,0,0,0.06); padding: 1px 4px; border-radius: 3px; font-size: 12px; }
.msg pre { background: rgba(0,0,0,0.06); padding: 8px 10px; border-radius: 6px; overflow-x: auto; margin: 6px 0; }
.msg pre code { background: none; padding: 0; }
.msg table { border-collapse: collapse; width: 100%; margin: 6px 0; font-size: 12px; }
.msg table th, .msg table td { padding: 4px 8px; border: 1px solid var(--kk-gray-200); text-align: left; }
.msg table th { background: var(--kk-gray-50); font-weight: 600; }
.msg blockquote { border-left: 3px solid var(--kk-gray-300); margin: 6px 0; padding: 2px 10px; color: var(--kk-gray-500); }
.chat-input { padding: 12px 16px; border-top: 1px solid var(--kk-gray-200); background: var(--kk-cream); display: flex; align-items: center; gap: 8px; }
.chat-input input { flex: 1; padding: 10px 16px; border: 1px solid var(--kk-gray-200); border-radius: 24px; font-size: 13px; font-family: inherit; background: white; outline: none; }
.chat-input input:focus { border-color: var(--kk-dark-red); box-shadow: 0 0 0 2px rgba(181,32,31,0.15); }
.chat-input button { width: 40px; height: 40px; border-radius: 50%; background: var(--kk-charcoal); color: white; border: none; cursor: pointer; display: flex; align-items: center; justify-content: center; transition: background 0.2s; flex-shrink: 0; }
.chat-input button:hover:not(:disabled) { background: var(--kk-dark-red); }
.chat-input button:disabled { opacity: 0.4; cursor: not-allowed; }
.chat-disclaimer { text-align: center; font-size: 11px; color: var(--kk-gray-400); padding: 6px 0 12px; }

/* === Results panel (mockup: tabs + white bg) === */
.results-panel { width: 60%; display: flex; flex-direction: column; overflow: hidden; min-height: 0; }
.results-content { flex: 1; overflow-y: auto; padding: 20px 8px; background: var(--kk-gray-50); border-radius: 0 0 8px 8px; }

/* === Component cards (mockup style) === */
.comp-card { background: white; border: 1px solid var(--kk-gray-200); border-radius: 8px; overflow: hidden; margin-bottom: 16px; }
.comp-card-header { padding: 12px 16px; background: var(--kk-gray-50); border-bottom: 1px solid var(--kk-gray-200); }
.comp-card-header h3 { font-size: 14px; font-weight: 600; color: var(--kk-charcoal); }
.comp-table { width: 100%; border-collapse: collapse; font-size: 13px; }
.comp-table th { padding: 8px 12px; text-align: left; font-weight: 500; color: var(--kk-gray-500); font-size: 12px; border-bottom: 1px solid var(--kk-gray-200); }
.comp-table td { padding: 10px 12px; border-bottom: 1px solid var(--kk-gray-100); }
.comp-table tr:last-child td { border-bottom: none; }
.alt-row { cursor: pointer; transition: background 0.15s; }
.alt-row:hover { background: var(--kk-gray-50); }
.alt-row.selected { background: var(--kk-gold-light) !important; }
.alt-row input[type=radio] { accent-color: var(--kk-charcoal); }
/* Per-component usage_context — subtle callout under component name in tables */
.usage-context { margin-top: 6px; padding: 8px 12px 8px 14px; background: var(--kk-gray-50); border-left: 2px solid var(--kk-gray-300); border-radius: 0 3px 3px 0; font-size: 12.5px; line-height: 1.5; color: var(--kk-gray-500); font-style: normal; }
.usage-context-label { display: block; font-size: 9.5px; font-weight: 700; letter-spacing: 1.3px; text-transform: uppercase; color: var(--kk-gray-400); margin: 0 0 3px; }

/* === Needs analysis card (editorial pairing) === */
.needs-card { background: white; border: 1px solid var(--kk-gray-200); border-radius: 6px; overflow: hidden; margin-bottom: 16px; }
.needs-card-head { padding: 14px 20px 12px; border-bottom: 1px solid var(--kk-gray-100); display: flex; align-items: baseline; justify-content: space-between; gap: 16px; flex-wrap: wrap; }
.needs-card-title { font-size: 15px; font-weight: 500; color: var(--kk-charcoal); }
.needs-card-sub { font-size: 11.5px; color: var(--kk-gray-500); font-style: italic; }
.needs-body { padding: 20px 24px 22px; }
.needs-empty { padding: 16px 20px; font-size: 12px; color: var(--kk-gray-500); }

/* Voice blocks */
.voice-block { position: relative; padding: 4px 0 4px 22px; }
.voice-block + .voice-block { margin-top: 0; }
.voice-block::before { content: ''; position: absolute; left: 0; top: 6px; bottom: 6px; width: 3px; border-radius: 2px; }
.voice-user::before { background: var(--kk-gray-300); }
.voice-aida::before { background: var(--kk-red); }
.voice-label { font-size: 10.5px; font-weight: 700; letter-spacing: 1.5px; text-transform: uppercase; margin-bottom: 6px; display: flex; align-items: center; gap: 8px; }
.voice-user .voice-label { color: var(--kk-gray-500); }
.voice-aida .voice-label { color: var(--kk-red); }
.voice-label .dot { display: inline-block; width: 6px; height: 6px; border-radius: 50%; background: currentColor; opacity: 0.7; }
.voice-text { font-size: 14.5px; line-height: 1.6; color: var(--kk-charcoal); }
.voice-user .voice-text { color: #5a5854; }
.voice-text em.empty { color: var(--kk-gray-400); }

/* Transition between user and aida */
.voice-transition { margin: 12px 0 12px 22px; font-size: 11.5px; color: var(--kk-gray-400); display: flex; align-items: center; gap: 8px; letter-spacing: 0.3px; }
.voice-transition::before { content: ''; height: 16px; border-left: 1.5px dashed var(--kk-gray-300); margin-left: -22px; width: 22px; }

/* Inferens edit affordance */
.voice-aida { position: relative; }
.voice-aida-actions { position: absolute; top: 0; right: 0; }
.voice-aida-edit { background: none; border: 1px solid var(--kk-gray-200); border-radius: 100px; padding: 4px 11px 4px 9px; font-size: 11px; color: var(--kk-gray-500); cursor: pointer; display: inline-flex; align-items: center; gap: 6px; font-family: inherit; transition: all 0.15s; }
.voice-aida-edit:hover { background: #FDF7F7; border-color: var(--kk-red); color: var(--kk-red); }
.voice-aida-edit svg { width: 11px; height: 11px; }
.voice-aida.is-editing .voice-text { display: none; }
.voice-aida.is-editing .voice-aida-edit { display: none; }
.voice-aida-textarea { display: none; width: 100%; min-height: 140px; border: 1.5px solid var(--kk-red); border-radius: 4px; padding: 12px 14px; font-family: inherit; font-size: 14.5px; line-height: 1.6; color: var(--kk-charcoal); background: white; resize: vertical; box-sizing: border-box; }
.voice-aida.is-editing .voice-aida-textarea { display: block; }
.voice-aida-textarea:focus { outline: none; box-shadow: 0 0 0 3px rgba(181, 32, 31, 0.15); }
.voice-aida-edit-actions { display: none; gap: 8px; margin-top: 10px; justify-content: flex-end; }
.voice-aida.is-editing .voice-aida-edit-actions { display: flex; }
.btn-na-cancel { background: none; border: 1px solid var(--kk-gray-300); color: var(--kk-gray-500); padding: 5px 14px; font-size: 11.5px; border-radius: 3px; cursor: pointer; font-family: inherit; }
.btn-na-save { background: var(--kk-charcoal); border: 1px solid var(--kk-charcoal); color: white; padding: 5px 14px; font-size: 11.5px; border-radius: 3px; cursor: pointer; font-family: inherit; }
.btn-na-save:hover { background: #2a2a2a; }

/* Meta blocks (assumptions + would_clarify) */
.needs-meta-row { margin-top: 22px; padding-top: 16px; border-top: 1px solid var(--kk-gray-100); display: grid; grid-template-columns: 1fr 1fr; gap: 24px; }
.needs-meta-label { font-size: 10.5px; font-weight: 700; letter-spacing: 1.3px; text-transform: uppercase; color: var(--kk-gray-400); margin-bottom: 6px; }
.needs-meta-list { list-style: none; margin: 0; padding: 0; font-size: 12.5px; line-height: 1.55; color: var(--kk-gray-500); }
.needs-meta-list li { position: relative; padding: 3px 0 3px 18px; }
.needs-meta-list li::before { position: absolute; left: 0; top: 3px; }
.needs-meta-assumptions li::before { content: '·'; font-size: 18px; line-height: 1; color: var(--kk-gray-300); }
.needs-meta-clarify li::before { content: '?'; font-style: italic; color: var(--kk-red); opacity: 0.55; }
@media (max-width: 640px) {
  .needs-meta-row { grid-template-columns: 1fr; gap: 16px; }
  .needs-card-head { flex-direction: column; align-items: flex-start; gap: 4px; }
  .voice-aida-actions { position: static; margin-bottom: 6px; }
}
.type-badge { display: inline-block; padding: 2px 8px; border-radius: 12px; font-size: 11px; font-weight: 600; }
.type-baseline { background: var(--kk-gray-100); color: var(--kk-charcoal); }
.type-reuse { background: var(--kk-gold-light); color: #7A6000; }
.type-optimized { background: #FDE8D0; color: var(--kk-red-orange); }

/* === Summary cards === */
.summary { display: grid; grid-template-columns: repeat(3, 1fr); gap: 12px; margin: 16px 0; }
.card { background: var(--kk-gray-50); border: 1px solid var(--kk-gray-200); border-radius: 8px; padding: 16px; }
.card .card-title { font-size: 11px; font-weight: 600; color: var(--kk-gray-500); text-transform: uppercase; letter-spacing: 0.5px; }
.card .value { font-size: 24px; font-weight: 700; color: var(--kk-charcoal); margin-top: 4px; }
.card .sublabel { font-size: 12px; color: var(--kk-gray-500); }
.card.saving .value { color: var(--green-saving); }

/* === Source badges === */
.source-badge { display: inline-block; padding: 2px 6px; border-radius: 4px; font-size: 10px; font-weight: 600; margin-right: 3px; }
.source-verified { background: #F0E0E0; color: var(--kk-burgundy); }
.source-aggregate { background: #FFE9D6; color: #7A4810; }
.source-estimate { background: var(--kk-gold-light); color: #8B6914; }
.source-legend { display: flex; gap: 16px; margin: 4px 0 12px; font-size: 12px; color: var(--kk-gray-500); }
.method-label { margin: 4px 0 8px; font-size: 11px; color: var(--kk-gray-500); font-style: italic; }

/* === Buttons === */
.btn { padding: 10px 20px; background: var(--kk-dark-red); color: white; border: none; border-radius: 8px; cursor: pointer; font-size: 13px; font-weight: 600; font-family: inherit; margin-top: 12px; transition: background 0.2s; }
.btn:hover { background: var(--kk-burgundy); }
.btn:disabled { opacity: 0.4; cursor: not-allowed; }
.btn-secondary { background: var(--kk-gray-100); color: var(--kk-charcoal); }
.btn-secondary:hover { background: var(--kk-gray-200); }

.section-title { font-size: 15px; font-weight: 600; margin: 16px 0 6px; color: var(--kk-charcoal); }
.report-area { background: white; border: 1px solid var(--kk-gray-200); border-radius: 8px; padding: 20px; margin-top: 16px; font-size: 13px; line-height: 1.6; max-height: 500px; overflow-y: auto; }
.report-area h1 { font-size: 20px; font-weight: 700; margin: 0 0 12px; color: var(--kk-charcoal); border-bottom: 2px solid var(--kk-gray-200); padding-bottom: 6px; }
.report-area h2 { font-size: 16px; font-weight: 600; margin: 16px 0 8px; color: var(--kk-charcoal); }
.report-area h3 { font-size: 14px; font-weight: 600; margin: 12px 0 6px; color: var(--kk-charcoal); }
.report-area p { margin: 0 0 10px; }
.report-area ul, .report-area ol { margin: 6px 0 10px; padding-left: 24px; }
.report-area li { margin-bottom: 4px; }
.report-area table { border-collapse: collapse; width: 100%; margin: 10px 0; font-size: 12px; }
.report-area table th, .report-area table td { padding: 6px 10px; border: 1px solid var(--kk-gray-200); text-align: left; }
.report-area table th { background: var(--kk-gray-50); font-weight: 600; font-size: 11px; color: var(--kk-gray-500); }
.report-area strong { font-weight: 600; }
.report-area blockquote { border-left: 3px solid var(--kk-gold); margin: 8px 0; padding: 4px 12px; background: var(--kk-cream); color: var(--kk-gray-500); font-style: italic; }
.report-area hr { border: none; border-top: 1px solid var(--kk-gray-200); margin: 16px 0; }

/* === Footer (karlstad.se: warm cream) === */
.footer { background: var(--kk-cream); color: var(--kk-gray-500); height: 36px; display: flex; align-items: center; justify-content: center; font-size: 11px; flex-shrink: 0; border-top: 1px solid var(--kk-gray-200); }

/* === Scrollbar === */
::-webkit-scrollbar { width: 8px; }
::-webkit-scrollbar-track { background: transparent; }
::-webkit-scrollbar-thumb { background: var(--kk-gray-300); border-radius: 8px; }
::-webkit-scrollbar-thumb:hover { background: var(--kk-gray-400); }
html { scrollbar-width: thin; scrollbar-color: #d4d4d4 transparent; }

.empty-state { color: var(--kk-gray-400); text-align: center; margin-top: 80px; }
.empty-state p { font-size: 14px; }

/* === Auth overlay === */
#authOverlay { display: flex; align-items: center; justify-content: center; flex: 1; background: var(--kk-warm-bg); }
#authOverlay .login-box { background: white; border-radius: 12px; padding: 40px; width: 360px; box-shadow: 0 4px 24px rgba(0,0,0,0.08); border-top: 3px solid var(--kk-gold-light); }
#authOverlay .login-box h1 { font-size: 24px; color: var(--kk-charcoal); margin-bottom: 8px; }
#authOverlay .login-box p { font-size: 13px; color: var(--kk-gray-400); margin-bottom: 24px; }
#authOverlay .login-box input { width: 100%; padding: 12px 16px; border: 1px solid var(--kk-gray-200); border-radius: 8px; font-size: 14px; font-family: inherit; outline: none; margin-bottom: 8px; }
#authOverlay .login-box input:focus { border-color: var(--kk-dark-red); box-shadow: 0 0 0 2px rgba(181,32,31,0.15); }
#authOverlay .login-box button { width: 100%; padding: 12px; background: var(--kk-charcoal); color: white; border: none; border-radius: 8px; font-size: 14px; font-weight: 600; cursor: pointer; margin-top: 4px; font-family: inherit; }
#authOverlay .login-box button:hover { background: var(--kk-dark-red); }
#authOverlay .login-box button:disabled { opacity: 0.4; cursor: not-allowed; }
#authOverlay .error { color: var(--kk-dark-red); font-size: 12px; margin: 4px 0; }
#appContainer { display: flex; flex-direction: column; flex: 1; min-height: 0; }

/* === Dropdown menus === */
.project-btn { background: none; border: none; color: var(--kk-gray-500); font-size: 14px; cursor: pointer; display: flex; align-items: center; gap: 6px; padding: 6px 12px; border-radius: 6px; font-family: inherit; }
.project-btn:hover { background: var(--kk-gray-100); color: var(--kk-charcoal); }
.user-btn { background: none; border: none; color: var(--kk-gray-400); cursor: pointer; padding: 6px; border-radius: 50%; display: flex; align-items: center; }
.user-btn:hover { background: var(--kk-gray-100); color: var(--kk-charcoal); }
.dropdown-menu { position: absolute; top: calc(100% + 4px); background: white; border: 1px solid var(--kk-gray-200); border-radius: 8px; box-shadow: 0 4px 16px rgba(0,0,0,0.12); min-width: 220px; z-index: 100; padding: 4px 0; }
.dropdown-right { right: 0; }
.dropdown-header { padding: 8px 16px; font-size: 11px; font-weight: 600; color: var(--kk-gray-400); text-transform: uppercase; }
.dropdown-divider { border-top: 1px solid var(--kk-gray-200); margin: 4px 0; }
.dropdown-item { display: flex; align-items: center; gap: 8px; width: 100%; padding: 8px 16px; border: none; background: none; font-size: 13px; color: var(--kk-charcoal); cursor: pointer; font-family: inherit; text-align: left; }
.dropdown-item:hover { background: var(--kk-gray-50); }
.dropdown-item.active { background: var(--kk-gold-light); }

/* === Results tabs === */
.results-tabs { display: flex; border-bottom: 2px solid var(--kk-gray-200); flex-shrink: 0; background: white; border-radius: 8px 8px 0 0; }
.tab { padding: 10px 20px; background: none; border: none; font-size: 13px; font-weight: 500; color: var(--kk-gray-400); cursor: pointer; border-bottom: 2px solid transparent; margin-bottom: -2px; font-family: inherit; transition: all 0.2s; }
.tab:hover:not(:disabled) { color: var(--kk-charcoal); }
.tab.active { color: var(--kk-charcoal); border-bottom-color: var(--kk-dark-red); font-weight: 600; }
.tab:disabled { opacity: 0.35; cursor: not-allowed; }

/* === Confirm actions in chat (legacy, kept for history rendering) === */
.confirm-actions { display: none; }
.confirm-hint { display: none; }

/* === Sticky confirm bar above chat input === */
.confirm-bar { display: none; padding: 10px 16px; background: var(--kk-cream); border-top: 1px solid var(--kk-gray-200); align-items: center; gap: 10px; }
.confirm-bar.visible { display: flex; }
.confirm-bar .confirm-bar-text { flex: 1; font-size: 12px; color: var(--kk-gray-500); }
.confirm-bar .btn-confirm-sticky { padding: 8px 20px; background: var(--kk-charcoal); color: white; border: none; border-radius: 20px; font-size: 12px; font-weight: 600; cursor: pointer; font-family: inherit; transition: background 0.2s; }
.confirm-bar .btn-confirm-sticky:hover:not(:disabled) { background: var(--kk-dark-red); }
.confirm-bar .btn-confirm-sticky:disabled { opacity: 0.4; cursor: not-allowed; }

/* === Typing indicator (Feature 1) === */
.typing-indicator { display: flex; align-items: center; gap: 5px; padding: 10px 14px; }
.typing-dot { width: 7px; height: 7px; border-radius: 50%; background: var(--kk-gray-400); animation: typingBounce 1.2s ease-in-out infinite; }
.typing-dot:nth-child(2) { animation-delay: 0.2s; }
.typing-dot:nth-child(3) { animation-delay: 0.4s; }
@keyframes typingBounce { 0%, 60%, 100% { transform: translateY(0); opacity: 0.4; } 30% { transform: translateY(-6px); opacity: 1; } }
.elapsed-time { font-size: 11px; color: var(--kk-gray-400); margin-left: 4px; }
.typing-text { font-size: 12px; color: var(--kk-gray-500); margin-left: 8px; font-style: italic; }
.action-btn { padding: 6px 14px; background: var(--kk-charcoal); color: white; border: none; border-radius: 16px; font-size: 12px; font-weight: 600; cursor: pointer; font-family: inherit; transition: background 0.2s; }
.action-btn:hover:not(:disabled) { background: var(--kk-dark-red); }
.action-btn:disabled { opacity: 0.5; cursor: not-allowed; }
.project-rename-input { background: transparent; border: 1px solid var(--kk-gray-300); border-radius: 4px; padding: 2px 6px; font-size: inherit; font-family: inherit; color: inherit; outline: none; min-width: 120px; }
.project-rename-input:focus { border-color: var(--kk-charcoal); }

/* === Reasoning expander (Feature 2) === */
.reasoning-toggle { background: none; border: none; color: var(--kk-gray-400); font-size: 11px; cursor: pointer; padding: 0; font-family: inherit; text-decoration: underline; white-space: nowrap; }
.reasoning-toggle:hover { color: var(--kk-charcoal); }
.reasoning-row td { padding: 4px 12px 8px 44px; font-size: 12px; color: var(--kk-gray-500); line-height: 1.5; background: var(--kk-gray-50); border-bottom: 1px solid var(--kk-gray-100); }

/* === Modal (Feature 5) === */
.modal-backdrop { position: fixed; inset: 0; background: rgba(0,0,0,0.4); z-index: 200; display: flex; align-items: center; justify-content: center; }
.modal-box { background: white; border-radius: 12px; padding: 32px; max-width: 560px; width: 90%; max-height: 80vh; overflow-y: auto; position: relative; }
.modal-box h2 { font-size: 18px; font-weight: 700; color: var(--kk-charcoal); margin-bottom: 16px; }
.modal-box p, .modal-box li { font-size: 13px; line-height: 1.6; color: var(--kk-charcoal); }
.modal-box ul { padding-left: 20px; margin: 8px 0; }
.modal-box section { margin-bottom: 20px; }
.modal-box h3 { font-size: 14px; font-weight: 600; color: var(--kk-charcoal); margin-bottom: 6px; }
.modal-close { position: absolute; top: 16px; right: 16px; background: none; border: none; cursor: pointer; color: var(--kk-gray-400); font-size: 20px; }
.modal-close:hover { color: var(--kk-charcoal); }

/* === Step-back navigation (Feature 10) === */
.step-circle.done { cursor: pointer; }
.step-circle.done:hover { background: var(--kk-dark-red); border-color: var(--kk-dark-red); transform: scale(1.1); transition: all 0.2s; }

/* === Responsive (Feature 9) === */
@media (max-width: 768px) {
  .main { flex-direction: column; overflow-y: auto; overflow-x: hidden; padding: 0 12px; gap: 12px; }
  .chat-panel { width: 100%; min-height: 300px; max-height: 50vh; }
  .results-panel { width: 100%; }
  .progress-bar { padding: 12px 16px 8px; }
  .step-label { display: none; }
  .progress-track { gap: 0; }
  .step-circle { width: 26px; height: 26px; font-size: 12px; }
  .topbar { padding: 0 12px; }
  .topbar-center { max-width: 120px; overflow: hidden; text-overflow: ellipsis; white-space: nowrap; }
  .summary { grid-template-columns: repeat(2, 1fr); }
  .card .value { font-size: 18px; }
  .comp-table { display: block; overflow-x: auto; -webkit-overflow-scrolling: touch; }
  body { height: auto; min-height: 100vh; }
  .results-content { max-height: none; }
  .footer { padding: 8px 12px; }
}
@media (max-width: 480px) {
  .summary { grid-template-columns: 1fr; }
  .chat-panel { max-height: 45vh; }
  .results-tabs { overflow-x: auto; }
  .tab { padding: 10px 12px; font-size: 12px; white-space: nowrap; }
  .topbar-center { display: none; }
}
</style>
</head>
<body>

{% if has_supabase %}
<!-- Auth overlay -->
<div id="authOverlay">
  <div class="login-box">
    <h1>AIda</h1>
    <p>Klimatkalkyl och beslutsstöd för ombyggnationer</p>
    <input type="email" id="authEmail" placeholder="E-post" autofocus>
    <input type="password" id="authPassword" placeholder="Lösenord" onkeydown="if(event.key==='Enter')handleAuth()">
    <div id="authError" class="error" style="display:none"></div>
    <button onclick="handleAuth()" id="authSubmitBtn">Logga in</button>
    <div style="text-align:center;margin-top:12px;font-size:13px;color:var(--kk-gray-400)">
      <span id="authToggleText">Inget konto?</span>
      <a href="#" onclick="toggleAuthMode(event)" id="authToggleLink" style="color:var(--kk-dark-red)">Skapa konto</a>
    </div>
  </div>
</div>
<div id="appContainer">
{% endif %}

<!-- Top bar -->
<div class="topbar">
  <div class="topbar-logo">
    <svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" stroke-linecap="round"><circle cx="12" cy="12" r="5"/><path d="M12 1v3M12 20v3M1 12h3M20 12h3M4.2 4.2l2.1 2.1M17.7 17.7l2.1 2.1M4.2 19.8l2.1-2.1M17.7 6.3l2.1-2.1"/></svg>
    <span>AIda</span>
  </div>
  {% if has_supabase %}
  <div class="topbar-center" id="projectDropdown" style="position:relative">
    <button class="project-btn" onclick="toggleProjectMenu()" id="projectBtn">
      <span id="projectName">Nytt projekt</span>
      <svg width="12" height="12" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><polyline points="6 9 12 15 18 9"/></svg>
    </button>
    <div class="dropdown-menu" id="projectMenu" style="display:none;left:50%;transform:translateX(-50%)">
      <div class="dropdown-header">Senaste projekt</div>
      <div id="projectList"></div>
      <div class="dropdown-divider"></div>
      <button class="dropdown-item" onclick="startRenameProject()">Byt namn på projektet</button>
      <button class="dropdown-item" onclick="createNewProject()">+ Skapa nytt projekt</button>
    </div>
  </div>
  <div class="topbar-right" id="userDropdown" style="position:relative;display:flex;align-items:center;gap:12px">
    <a href="#" onclick="openAbout();return false" style="color:var(--kk-gray-500);text-decoration:none;font-size:12px">Om verktyget</a>
    <span id="saveIndicator" style="font-size:11px;color:var(--kk-gray-400);display:none"></span>
    <button class="user-btn" onclick="toggleUserMenu()">
      <svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5"><path d="M20 21v-2a4 4 0 0 0-4-4H8a4 4 0 0 0-4 4v2"/><circle cx="12" cy="7" r="4"/></svg>
    </button>
    <div class="dropdown-menu dropdown-right" id="userMenu" style="display:none">
      <div class="dropdown-header" id="userEmail"></div>
      <div class="dropdown-divider"></div>
      <button class="dropdown-item" onclick="handleLogout()">Logga ut</button>
    </div>
  </div>
  {% else %}
  <div class="topbar-center"></div>
  <div class="topbar-right" style="display:flex;align-items:center;gap:12px"><a href="#" onclick="openAbout();return false" style="color:var(--kk-gray-500);text-decoration:none;font-size:12px">Om verktyget</a><span id="saveIndicator" style="font-size:11px;color:var(--kk-gray-400);display:none"></span><span style="font-size:12px;color:var(--kk-gray-400)">Prototyp</span></div>
  {% endif %}
</div>

<!-- Progress tracker -->
<div class="progress-bar">
  <div class="progress-track">
    <div class="progress-line"><div class="progress-fill" id="progressFill" style="width:0%"></div></div>
    <div class="step-item" data-step="planering">
      <div class="step-circle" id="sc-planering">1</div>
      <div class="step-label" id="sl-planering">Projektbeskrivning</div>
    </div>
    <div class="step-item" data-step="baslinje">
      <div class="step-circle" id="sc-baslinje">2</div>
      <div class="step-label" id="sl-baslinje">Baslinje</div>
    </div>
    <div class="step-item" data-step="aterbruk">
      <div class="step-circle" id="sc-aterbruk">3</div>
      <div class="step-label" id="sl-aterbruk">&#xC5;terbruk</div>
    </div>
    <div class="step-item" data-step="nyproduktion">
      <div class="step-circle" id="sc-nyproduktion">4</div>
      <div class="step-label" id="sl-nyproduktion">Nyproduktion</div>
    </div>
    <div class="step-item" data-step="sammanstallning">
      <div class="step-circle" id="sc-sammanstallning">5</div>
      <div class="step-label" id="sl-sammanstallning">Sammanst&#xE4;llning</div>
    </div>
    <div class="step-item" data-step="uppfoljning">
      <div class="step-circle" id="sc-uppfoljning">6</div>
      <div class="step-label" id="sl-uppfoljning">Uppf&#xF6;ljning</div>
    </div>
  </div>
</div>

<!-- Main content -->
<div class="main">
  <!-- Chat panel -->
  <div class="chat-panel">
    <div class="chat-container">
      <div class="chat-header">
        <h2>AIda</h2>
      </div>
      <div class="messages" id="messages">
        <div class="msg bot">Hej! Beskriv ditt projekt. Ange byggnadstyp, byggnadsår, ungefärlig yta och vilka behoven är.</div>
      </div>
      <div class="confirm-bar" id="confirmBar">
        <span class="confirm-bar-text" id="confirmBarText"></span>
        <button class="btn-confirm-sticky" id="confirmBarBtn" onclick="confirmStep()"></button>
      </div>
      <div class="chat-input">
        <input id="userInput" type="text" placeholder="Skriv ditt meddelande..." onkeydown="if(event.key==='Enter')sendMessage()">
        <button id="sendBtn" onclick="sendMessage()" aria-label="Skicka">
          <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><line x1="22" y1="2" x2="11" y2="13"/><polygon points="22 2 15 22 11 13 2 9 22 2"/></svg>
        </button>
      </div>
    </div>
    <div class="chat-disclaimer"></div>
  </div>

  <!-- Results panel -->
  <div class="results-panel" id="results">
    <div class="results-tabs" id="resultTabs" style="display:none">
      <button class="tab" id="tab-projekt" onclick="switchTab('projekt')" disabled>Projekt</button>
      <button class="tab" id="tab-baslinje" onclick="switchTab('baslinje')" disabled>Baslinje</button>
      <button class="tab" id="tab-alternativ" onclick="switchTab('alternativ')" disabled>Alternativ</button>
      <button class="tab" id="tab-rapport" onclick="switchTab('rapport')" disabled>Rapport</button>
    </div>
    <div class="results-content" id="resultContent">
      <div class="empty-state">
        <p>Beskriv ditt projekt i chatten till vänster för att börja.</p>
      </div>
    </div>
  </div>
</div>

<!-- Footer -->
<div class="footer" style="display:flex;justify-content:center;align-items:center;gap:8px">
  <span style="font-size:11px;color:var(--kk-gray-400)">AIda kan g&#xF6;ra misstag. Kontrollera viktig information.</span>
</div>

<!-- About modal (Feature 5) -->
<div id="aboutModal" class="modal-backdrop" style="display:none" onclick="if(event.target===this)closeAbout()">
  <div class="modal-box">
    <button class="modal-close" onclick="closeAbout()" aria-label="St&#xE4;ng">&#x2715;</button>
    <h2>Om AIda</h2>
    <section>
      <h3>Vad &#xE4;r AIda?</h3>
      <p>AIda &#xE4;r ett AI-drivet beslutsst&#xF6;d f&#xF6;r klimatber&#xE4;kning vid ombyggnation av kommunala fastigheter. Verktyget utvecklas inom Klimatneutrala Karlstad 2030, finansierat av Vinnova, Energimyndigheten och Formas inom ramen f&#xF6;r strategiska innovationsprogrammet Viable Cities.</p>
    </section>
    <section>
      <h3>Datak&#xE4;llor</h3>
      <ul>
        <li><strong>Klimatdata:</strong> Boverkets klimatdatabas, NollCO2-metoden</li>
        <li><strong>Alternativ:</strong> Environdec EPD-databas (verifierade produktdeklarationer)</li>
        <li><strong>Priser:</strong> AI-driven webbs&#xF6;kning mot svenska bygghandlare</li>
        <li><strong>&#xC5;terbruk:</strong> Palats (Karlstads kommuns &#xE5;terbruksplattform)</li>
      </ul>
    </section>
    <section>
      <h3>Metod</h3>
      <p>AIda j&#xE4;mf&#xF6;r konventionella materialval (baslinje) mot klimatoptimerade alternativ med hj&#xE4;lp av verifierade EPD:er. Ber&#xE4;kningarna avser produktskedet (A1-A3) om inget annat anges.</p>
    </section>
    <section>
      <h3>Begr&#xE4;nsningar</h3>
      <ul>
        <li>Resultaten &#xE4;r ett underlag f&#xF6;r beslut, inte ett slutgiltigt klimatbokslut</li>
        <li>Kostnadsuppskattningar baseras p&#xE5; webbs&#xF6;kning &#x2014; inh&#xE4;mta offerter f&#xF6;r exakta v&#xE4;rden</li>
        <li>AI kan g&#xF6;ra fel &#x2014; kontrollera k&#xE4;llh&#xE4;nvisningar vid viktiga beslut</li>
      </ul>
    </section>
    <section>
      <h3>Kontakt</h3>
      <p>Henric Barkman, <a href="mailto:henric.barkman@karlstad.se" style="color:var(--kk-blue)">henric.barkman@karlstad.se</a></p>
    </section>
  </div>
</div>

{% if has_supabase %}</div><!-- /appContainer -->{% endif %}

<script>
// Configure marked
if (typeof marked !== 'undefined') {
  marked.setOptions({ breaks: true, gfm: true });
}

let _step = 'idle';
let state = {
  project: null, baseline: null, alternatives: null,
  selections: {}, pendingDesc: null, reportMarkdown: null,
  chatHistory: [],
  get step() { return _step; },
  set step(v) { _step = v; updatePlaceholder(); },
};
let activeTab = null;

// Dynamic placeholder (Feature 4)
const STEP_PLACEHOLDERS = {
  idle: 'Beskriv ditt ombyggnadsprojekt...',
  intake_done: 'Korrigera eller bekr\u00e4fta...',
  baseline_done: 'Diskutera, korrigera eller bekr\u00e4fta...',
  alternatives_done: 'Diskutera, korrigera eller generera rapport...',
  report_done: 'Diskutera eller korrigera analysen...',
};
function updatePlaceholder() {
  document.getElementById('userInput').placeholder = STEP_PLACEHOLDERS[state.step] || 'Skriv ditt meddelande...';
  updateConfirmBar();
}

// Sticky confirm bar
const CONFIRM_BAR_CONFIG = {
  intake_done: { text: 'Projektbeskrivning klar.', btn: 'Ber\u00e4kna baslinje \u2192' },
  baseline_done: { text: 'Baslinjen \u00e4r klar.', btn: 'S\u00f6k alternativ \u2192' },
  alternatives_done: { text: 'Alternativ redo.', btn: 'Generera rapport \u2192' },
};
function updateConfirmBar() {
  const bar = document.getElementById('confirmBar');
  const cfg = CONFIRM_BAR_CONFIG[state.step];
  if (cfg) {
    document.getElementById('confirmBarText').textContent = cfg.text;
    const btn = document.getElementById('confirmBarBtn');
    btn.textContent = cfg.btn;
    btn.disabled = false;
    btn.style.opacity = '';
    bar.classList.add('visible');
  } else {
    bar.classList.remove('visible');
  }
}

function esc(s) { return String(s==null?'':s).replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;').replace(/"/g,'&quot;'); }

function renderMd(text) {
  text = text.replace(/^(\d+)\)\s/gm, '$1. ');
  let html;
  if (typeof marked !== 'undefined') html = marked.parse(text);
  else html = text.replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;')
    .replace(/\*\*(.+?)\*\*/g,'<strong>$1</strong>').replace(/\n/g,'<br>');
  return typeof DOMPurify !== 'undefined' ? DOMPurify.sanitize(html) : html;
}

let chatLog = []; // {text, cls, confirm?{btnLabel,hint}} — persisted via localStorage

function _chatStorageKey() { return 'aida_chat_' + (currentAnalysisId || 'new'); }

function _saveChatLog() {
  try { localStorage.setItem(_chatStorageKey(), JSON.stringify(chatLog)); } catch(e) {}
}

function addMsg(text, cls) {
  chatLog.push({text, cls});
  _saveChatLog();
  const d = document.createElement('div');
  d.className = 'msg ' + cls;
  if (cls === 'bot' || cls === 'system') { d.innerHTML = renderMd(text); }
  else { d.textContent = text; }
  document.getElementById('messages').appendChild(d);
  d.scrollIntoView({behavior:'smooth'});
}

function addConfirmMsg(text, btnLabel, hint) {
  chatLog.push({text, cls: 'bot', confirm: {btnLabel, hint}});
  _saveChatLog();
  const d = document.createElement('div');
  d.className = 'msg bot';
  d.innerHTML = renderMd(text) +
    '<div class="confirm-actions"><button class="btn-confirm" onclick="confirmStep()">' + btnLabel + '</button></div>' +
    '<div class="confirm-hint">' + hint + '</div>';
  document.getElementById('messages').appendChild(d);
  d.scrollIntoView({behavior:'smooth'});
}

function removeConfirmButtons() {
  document.querySelectorAll('.confirm-actions').forEach(el => {
    const btn = el.querySelector('.btn-confirm');
    if (btn) {
      btn.disabled = true;
      btn.style.cssText = 'background:var(--kk-gray-200);color:var(--kk-gray-400);cursor:default;pointer-events:none';
      btn.textContent = 'Bekr\u00e4ftad \u2713';
    }
    const hint = el.closest('.msg')?.querySelector('.confirm-hint');
    if (hint) hint.remove();
  });
}

function setProgressStep(name) {
  const order = ['planering','baslinje','aterbruk','nyproduktion','sammanstallning','uppfoljning'];
  const STEP_TAB = {planering:'projekt',baslinje:'baslinje',aterbruk:'alternativ',nyproduktion:'alternativ',sammanstallning:'alternativ',uppfoljning:'rapport'};
  const ni = order.indexOf(name);
  const pct = order.length > 1 ? (ni / (order.length - 1)) * 100 : 0;
  document.getElementById('progressFill').style.width = pct + '%';
  order.forEach((s, i) => {
    const circle = document.getElementById('sc-' + s);
    const label = document.getElementById('sl-' + s);
    const isDone = i < ni;
    circle.className = 'step-circle' + (isDone ? ' done' : i === ni ? ' active' : '');
    label.className = 'step-label' + (isDone ? ' done' : i === ni ? ' active' : '');
    if (isDone) {
      circle.innerHTML = '<svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="3" stroke-linecap="round" stroke-linejoin="round"><polyline points="20 6 9 17 4 12"/></svg>';
      const tab = STEP_TAB[s];
      circle.onclick = () => { if (tab) switchTab(tab); };
      circle.title = 'G\u00e5 till ' + label.textContent;
    } else {
      circle.textContent = i + 1;
      circle.onclick = null; circle.title = '';
    }
  });
}

let _loadingTimer = null;
let _loadingStart = null;
function setLoading(on) {
  document.getElementById('sendBtn').disabled = on;
  document.getElementById('userInput').disabled = on;
  // Always clean up previous loading state first (prevents duplicates)
  if (_loadingTimer) { clearInterval(_loadingTimer); _loadingTimer = null; }
  const prev = document.getElementById('typingBubble'); if (prev) prev.remove();
  if (on) {
    _loadingStart = Date.now();
    const el = document.createElement('div');
    el.className = 'msg bot'; el.id = 'typingBubble';
    el.innerHTML = '<div class="typing-indicator"><div class="typing-dot"></div><div class="typing-dot"></div><div class="typing-dot"></div><span class="typing-text" id="typingText">AIda jobbar...</span><span class="elapsed-time" id="elapsedTime"></span></div>';
    document.getElementById('messages').appendChild(el);
    el.scrollIntoView({behavior:'smooth'});
    _loadingTimer = setInterval(() => {
      const s = Math.floor((Date.now() - _loadingStart) / 1000);
      const t = document.getElementById('elapsedTime');
      if (t && s >= 3) t.textContent = s + 's';
      const tx = document.getElementById('typingText');
      if (tx && s >= 5 && !tx.dataset.long) {
        tx.textContent = 'AIda jobbar. Detta kan ta cirka 1-3 minuter.';
        tx.dataset.long = '1';
      }
    }, 1000);
  } else {
    updatePlaceholder();
  }
}

// === Tab system ===
function enableTab(name) {
  const tab = document.getElementById('tab-' + name);
  if (tab) tab.disabled = false;
  document.getElementById('resultTabs').style.display = 'flex';
}

function switchTab(name) {
  activeTab = name;
  document.querySelectorAll('.results-tabs .tab').forEach(t => t.classList.remove('active'));
  const tab = document.getElementById('tab-' + name);
  if (tab) tab.classList.add('active');
  // Render from state
  if (name === 'projekt' && state.project) renderProjektContent();
  else if (name === 'baslinje' && state.baseline) renderBaslinjeContent();
  else if (name === 'alternativ' && state.alternatives) renderAlternativContent();
  else if (name === 'rapport' && state.reportMarkdown) renderRapportContent();
}

// === Chat input ===
// Note: "kör" tas medvetet bort — det matchade "Kör omräkningen" och triggade
// oavsiktligt nästa steg. "kör vidare" fungerar fortfarande via "vidare".
const ADVANCE_RE = /\b(vidare|nästa|fortsätt|gå vidare|next|confirm|bekräfta)\b/i;
const ADVANCE_EXACT_RE = /^(ok|okej|ja)$/i;
const CORRECTION_RE = /\b(ändra|nej|fel|byt|korrigera|gör om|uppdatera|ta bort|lägg till|ändring|rätta|fixa|nytt? antal|inte \d|ska vara|stämmer inte|borde vara)\b/i;

// Build the context string we feed back into intake when the user makes a
// correction. Includes everything intake might otherwise re-ask about:
// projektnamn, byggnadstyp, area, komponenter, plus a tail of chat history so
// previously-answered clarifications (byggnadsår, krav, omfattning) are not
// re-asked. Reason: intake.py runs stateless on a single string description,
// so anything the model needs must travel in that string.
// Strip delimiter strings from message content so user or model text cannot
// impersonate the structural headers in the prompt below. Without this, a chat
// turn containing "Korrigering från användaren:" would inject a fake
// correction section into the next intake call.
function _scrubCtxDelimiters(s) {
  if (typeof s !== 'string') return '';
  return s
    .replace(/Korrigering från användaren:/gi, '[korrigering]')
    .replace(/Tidigare diskussion i sessionen:/gi, '[diskussion]')
    .replace(/^(Användare|AIda):/gim, '$1​:');  // zero-width space prevents role-line spoofing
}

function buildCorrectionContext(text) {
  // Guard: without core project fields we cannot build a useful ctx; pass the
  // raw correction text so intake at least sees the user's intent without
  // "undefined, undefined m2" garbage.
  if (!state.project || !state.project.building_type || state.project.area_bta == null) {
    return text;
  }
  const compSummary = (state.project.components || []).map(c => c.name + ' (' + c.quantity + ' ' + c.unit + ')').join(', ');
  let ctx = '';
  if (state.project.name) ctx += 'Projektnamn: ' + state.project.name + '. ';
  ctx += state.project.building_type + ', ' + state.project.area_bta + ' m2. Komponenter: ' + compSummary + '.';
  if (state.chatHistory && state.chatHistory.length > 0) {
    const tail = state.chatHistory.slice(-8).map(m => {
      const role = (m.role === 'user') ? 'Användare' : 'AIda';
      return role + ': ' + _scrubCtxDelimiters(m.content || '');
    }).join('\n');
    ctx += '\n\nTidigare diskussion i sessionen:\n' + tail;
  }
  ctx += '\n\nKorrigering från användaren: ' + _scrubCtxDelimiters(text);
  return ctx;
}

async function sendMessage() {
  const input = document.getElementById('userInput');
  const text = input.value.trim();
  if (!text) return;
  input.value = '';
  addMsg(text, 'user');
  setLoading(true);

  // Detect "advance to next step" intent at confirmation gates
  const wantsAdvance = ADVANCE_RE.test(text) || ADVANCE_EXACT_RE.test(text.trim());
  const wantsCorrection = CORRECTION_RE.test(text);

  switch (state.step) {
    case 'idle':
      if (state.pendingDesc) {
        await runIntake(state.pendingDesc + '\n\nFörtydligande: ' + text);
      } else {
        await runIntake(text);
      }
      break;
    case 'intake_done':
      if (wantsAdvance) {
        setLoading(false);
        confirmStep();
      } else {
        addMsg('Uppdaterar projektbeskrivning...', 'system');
        await runIntake(buildCorrectionContext(text));
      }
      break;
    case 'baseline_done':
      if (wantsAdvance) {
        setLoading(false);
        confirmStep();
      } else if (wantsCorrection) {
        // Re-run intake with correction, then auto-trigger baseline
        addMsg('Uppdaterar projektet och r\u00e4knar om baslinjen...', 'system');
        await runIntake(buildCorrectionContext(text));
        if (state.step === 'intake_done') {
          await runBaseline();
        }
      } else {
        await runChat(text);
      }
      break;
    case 'alternatives_done':
      if (wantsAdvance) {
        const allSel = state.alternatives && state.alternatives.components.every(c => state.selections[c.component_id]);
        if (!allSel) {
          const missing = state.alternatives.components.filter(c => !state.selections[c.component_id]).map(c => c.component_name);
          addMsg('Välj alternativ för alla komponenter först: ' + missing.join(', '), 'system');
          setLoading(false);
          break;
        }
        setLoading(false);
        generateReport();
      } else if (wantsCorrection) {
        // Re-run alternatives with user feedback
        addMsg('G\u00f6r om alternativs\u00f6kningen med dina kommentarer...', 'system');
        await runAlternatives(text);
      } else {
        await runChat(text);
      }
      break;
    case 'report_done':
      if (wantsCorrection) {
        // Re-run from intake with correction
        addMsg('G\u00f6r om analysen med dina kommentarer...', 'system');
        await runIntake(buildCorrectionContext(text));
        if (state.step === 'intake_done') {
          await runBaseline();
          if (state.step === 'baseline_done') {
            await runAlternatives();
          }
        }
      } else {
        await runChat(text);
      }
      break;
    default:
      setLoading(false);
  }
}

// === Confirm step ===
function confirmStep() {
  document.getElementById('confirmBarBtn').disabled = true;
  document.getElementById('confirmBarBtn').style.opacity = '0.5';
  if (state.step === 'intake_done') runBaseline();
  else if (state.step === 'baseline_done') runAlternatives();
  else if (state.step === 'alternatives_done') {
    const allSel = state.alternatives && state.alternatives.components.every(c => state.selections[c.component_id]);
    if (!allSel) {
      const missing = state.alternatives.components.filter(c => !state.selections[c.component_id]).map(c => c.component_name);
      addMsg('Välj alternativ för alla komponenter först: ' + missing.join(', '), 'system');
      document.getElementById('confirmBarBtn').disabled = false;
      document.getElementById('confirmBarBtn').style.opacity = '';
      return;
    }
    generateReport();
  }
}

// === Pipeline: Intake ===
async function runIntake(desc) {
  addMsg('Analyserar projektbeskrivning...', 'system');
  setProgressStep('planering');
  try {
    const r = await authFetch('/api/intake', {method:'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify({description: desc})});
    const d = await r.json();
    if (d.error) { addMsg('Fel: ' + d.error, 'system'); setLoading(false); return; }

    if (d.clarification_needed) {
      state.pendingDesc = desc;
      state.project = null;
      state.step = 'idle';
      if (d.components && d.components.length) {
        const list = d.components.map(c => '- ' + c.name).join('\n');
        addMsg('Hittade hittills:\n' + list, 'bot');
      }
      addMsg(d.clarification_needed, 'bot');
      setLoading(false);
      return;
    }

    state.pendingDesc = null;
    state.project = d;
    state.baseline = null;
    state.alternatives = null;
    state.selections = {};
    state.reportMarkdown = null;
    state.chatHistory = [];
    state.step = 'intake_done';
    if (HAS_SUPABASE) { document.getElementById('projectName').textContent = d.name || d.building_type || 'Nytt projekt'; }
    scheduleAutoSave();

    enableTab('projekt');
    switchTab('projekt');
    // Disable later tabs if re-running
    ['baslinje','alternativ','rapport'].forEach(t => { const el = document.getElementById('tab-'+t); if(el) el.disabled = true; });

    const compList = d.components.map(c => '- ' + c.name + ' (' + c.quantity + ' ' + c.unit + ')').join('\n');
    addConfirmMsg(
      '**' + d.building_type + '**, ' + d.area_bta + ' m\u00b2\n\n**Komponenter:**\n' + compList,
      'Bekr\u00e4fta och ber\u00e4kna baslinje \u2192',
      'Skriv i chatten om n\u00e5got inte st\u00e4mmer.'
    );
    setLoading(false);
  } catch(e) { addMsg('Fel: ' + e.message, 'system'); setLoading(false); }
}

// === Pipeline: Baseline ===
// Remove any lingering recompute action-row buttons. Used when we start a
// fresh baseline/alternatives run so an orphaned "S\u00f6k nya alternativ" button
// (rendered from a stale state earlier) does not stay clickable after we have
// already cleared state.alternatives.
function clearActionRows() {
  document.querySelectorAll('.msg.action-row').forEach(el => el.remove());
}

async function runBaseline() {
  addMsg('Ber\u00e4knar baslinje (NollCO2)...', 'system');
  setProgressStep('baslinje');
  setLoading(true);
  clearActionRows();
  try {
    const r = await authFetch('/api/baseline', {method:'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify({project: state.project})});
    const d = await r.json();
    if (d.error) {
      addMsg('Fel: ' + d.error, 'system');
      addConfirmMsg('Baslinjeberäkning misslyckades.', 'Försök igen \u2192', '');
      setLoading(false); return;
    }
    removeConfirmButtons();
    state.baseline = d;
    state.alternatives = null;
    state.selections = {};
    state.reportMarkdown = null;
    state.step = 'baseline_done';
    scheduleAutoSave();

    enableTab('baslinje');
    switchTab('baslinje');
    ['alternativ','rapport'].forEach(t => { const el = document.getElementById('tab-'+t); if(el) el.disabled = true; });

    const total = d.components.reduce((s,c) => s + c.co2e_kg, 0);
    addConfirmMsg(
      'Baslinje klar: **' + Math.round(total).toLocaleString('sv') + ' kg CO\u2082e** totalt f\u00f6r ' + d.components.length + ' komponenter.',
      'Bekr\u00e4fta och s\u00f6k alternativ \u2192',
      'Skriv i chatten om du vill korrigera n\u00e5got.'
    );
    setLoading(false);
  } catch(e) {
    addMsg('Fel: ' + e.message, 'system');
    addConfirmMsg('Baslinjeberäkning misslyckades.', 'Försök igen \u2192', '');
    setLoading(false);
  }
}

// === Pipeline: Alternatives ===
async function runAlternatives(userFeedback) {
  // Snapshot selections by component name before clearing (Feature 3)
  const prevSelByName = {};
  Object.values(state.selections).forEach(sel => { prevSelByName[sel.name] = sel.selected_alternative; });

  addMsg('S\u00f6ker alternativ...', 'system');
  setProgressStep('aterbruk');
  setLoading(true);
  clearActionRows();
  const subStepTimer = setTimeout(() => {
    setProgressStep('nyproduktion');
  }, 2000);
  try {
    const body = {project: state.project, baseline: state.baseline};
    if (userFeedback) body.user_feedback = userFeedback;
    const r = await authFetch('/api/alternatives', {method:'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify(body)});
    const d = await r.json();
    clearTimeout(subStepTimer);
    if (d.error) {
      clearTimeout(subStepTimer);
      addMsg('Fel: ' + d.error, 'system');
      addConfirmMsg('S\u00f6kning av alternativ misslyckades.', 'F\u00f6rs\u00f6k igen \u2192', '');
      setLoading(false); return;
    }
    removeConfirmButtons();
    state.alternatives = d;
    state.reportMarkdown = null;
    state.step = 'alternatives_done';

    // Restore previous selections by component name match (Feature 3)
    state.selections = {};
    if (Object.keys(prevSelByName).length > 0) {
      d.components.forEach(comp => {
        const prev = prevSelByName[comp.component_name];
        if (!prev) return;
        if (prev.name === 'Baslinje') {
          state.selections[comp.component_id] = {id:comp.component_id, name:comp.component_name, selected_alternative:{name:'Baslinje',co2e_kg:comp.baseline_co2e_kg,cost_sek:comp.baseline_cost_sek,source:'NollCO2'}, baseline_co2e_kg:comp.baseline_co2e_kg, baseline_cost_sek:comp.baseline_cost_sek};
        } else {
          const match = comp.alternatives.find(a => a.name === prev.name);
          if (match) state.selections[comp.component_id] = {id:comp.component_id, name:comp.component_name, selected_alternative:{name:match.name,co2e_kg:match.co2e_kg,cost_sek:match.cost_sek,source:match.source}, baseline_co2e_kg:comp.baseline_co2e_kg, baseline_cost_sek:comp.baseline_cost_sek};
        }
      });
    }
    scheduleAutoSave();
    setProgressStep('sammanstallning');

    enableTab('alternativ');
    switchTab('alternativ');
    document.getElementById('tab-rapport').disabled = true;

    const commentary = d.commentary || '';
    if (commentary) {
      addMsg(commentary, 'bot');
      addMsg('V\u00e4lj alternativ per komponent i resultatpanelen. Skriv i chatten om du vill ha fler f\u00f6rslag.', 'bot');
    } else {
      addMsg('Alternativ klara! V\u00e4lj per komponent i resultatpanelen.\n\nSkriv i chatten om du vill ha fler alternativ.', 'bot');
    }
    setLoading(false);
  } catch(e) {
    clearTimeout(subStepTimer);
    addMsg('Fel: ' + e.message, 'system');
    addConfirmMsg('S\u00f6kning av alternativ misslyckades.', 'F\u00f6rs\u00f6k igen \u2192', '');
    setLoading(false);
  }
}

// === Pipeline: Report ===
async function generateReport() {
  setProgressStep('uppfoljning');
  addMsg('Genererar rapport...', 'system');
  document.getElementById('reportBtn').disabled = true;
  setLoading(true);
  try {
    const sels = {components: Object.values(state.selections)};
    const r = await authFetch('/api/report', {method:'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify({project: state.project, selections: sels})});
    const d = await r.json();
    if (d.error) {
      addMsg('Fel: ' + d.error, 'system');
      addConfirmMsg('Rapportgenerering misslyckades.', 'Försök igen →', '');
      const rb = document.getElementById('reportBtn'); if (rb) rb.disabled = false;
      setLoading(false);
      return;
    }
    state.reportMarkdown = d.markdown;
    state.step = 'report_done';
    scheduleAutoSave();
    addMsg('Rapport klar!', 'bot');
    enableTab('rapport');
    switchTab('rapport');
    setLoading(false);
  } catch(e) {
    addMsg('Fel: ' + e.message, 'system');
    addConfirmMsg('Rapportgenerering misslyckades.', 'Försök igen →', '');
    const rb = document.getElementById('reportBtn'); if (rb) rb.disabled = false;
    setLoading(false);
  }
}

// === Conversational chat (agent with tool-use) ===
async function runChat(text) {
  setLoading(true);
  state.chatHistory.push({role:'user', content: text});
  try {
    const body = {
      message: text,
      history: state.chatHistory.slice(-10),
      project: state.project || null,
      baseline: state.baseline || null,
      alternatives: state.alternatives || null,
      selections: (state.selections && Object.keys(state.selections).length) ? state.selections : null,
    };
    const r = await authFetch('/api/chat', {method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify(body)});
    const d = await r.json();
    if (d.error) { addMsg('Fel: ' + d.error, 'system'); setLoading(false); return; }

    applyAgentStateUpdates(d.state_updates);

    state.chatHistory.push({role:'assistant', content: d.reply});
    addMsg(d.reply, 'bot');

    // Chat agent may have requested baseline/alternatives reruns. Execute them
    // sequentially so each action sees state from the previous one's merge.
    // Keep loading true throughout so the user cannot send a second message
    // mid-rerun and race the state.baseline / state.alternatives merge.
    const pendingActions = d.state_updates && d.state_updates.pending_actions;
    if (Array.isArray(pendingActions) && pendingActions.length > 0) {
      await processPendingActions(pendingActions);
    }
    setLoading(false);
  } catch(e) { addMsg('Fel: ' + e.message, 'system'); setLoading(false); }
}

// Execute reruns requested by the chat agent. Full reruns (empty component_ids)
// require an extra click; partial reruns run immediately since the user already
// initiated them via the chat correction. We pass orchestrated=true so the
// inner rerun functions do not toggle setLoading per action — runChat already
// holds setLoading(true) for the whole sequence to prevent the user from
// sending a second chat message that would race the state merges.
//
// Defense in depth against the prompt-only confirmation gate: if the agent
// emits explicit component_ids that cover every component in the project, we
// re-classify the action as full and route it to the confirmation button. The
// LLM cannot bypass the confirm gate by spelling out every id.
async function processPendingActions(actions) {
  const totalComponents = (state.project && Array.isArray(state.project.components))
    ? state.project.components.length : 0;
  const knownIds = new Set(
    (state.project && Array.isArray(state.project.components))
      ? state.project.components.map(c => c.id) : []
  );

  // Sort: rerun_baseline before rerun_alternatives for the same scope, so the
  // alternatives call sees the freshly recomputed baseline. The LLM can emit
  // them in any order and nothing in the schema enforces it.
  const ordered = [...actions].sort((a, b) => {
    if (a.type === b.type) return 0;
    if (a.type === 'rerun_baseline') return -1;
    if (b.type === 'rerun_baseline') return 1;
    return 0;
  });

  for (const action of ordered) {
    try {
      let cids = Array.isArray(action.component_ids) ? action.component_ids.filter(c => knownIds.has(c)) : [];
      // Empty after filter (unknown ids) and not originally full means we have
      // nothing actionable. Skip rather than treat as "all".
      const originallyEmpty = !Array.isArray(action.component_ids) || action.component_ids.length === 0;
      if (!originallyEmpty && cids.length === 0) {
        addMsg('Hoppar över ' + action.type + ': inga giltiga komponent-id.', 'system');
        continue;
      }
      // Re-classify as full when explicit list covers every component.
      const coversAll = totalComponents > 0 && cids.length === totalComponents;
      const isFull = originallyEmpty || coversAll;

      if (action.type === 'rerun_baseline') {
        if (isFull) {
          renderActionRow(
            'Bekräfta: räkna om hela baslinjen' + (action.reason ? ' (' + action.reason + ')' : ''),
            async () => { await runBaselineForComponents([], action.reason, false); },
          );
        } else {
          await runBaselineForComponents(cids, action.reason, true);
        }
      } else if (action.type === 'rerun_alternatives') {
        if (isFull) {
          renderActionRow(
            'Bekräfta: kör om alla alternativ' + (action.reason ? ' (' + action.reason + ')' : ''),
            async () => { await runAlternativesForComponents([], action.user_feedback, action.reason, false); },
          );
        } else {
          await runAlternativesForComponents(cids, action.user_feedback, action.reason, true);
        }
      } else {
        addMsg('Okänd åtgärd från chatten: ' + action.type, 'system');
      }
    } catch (e) {
      addMsg('Fel vid ' + action.type + ': ' + e.message, 'system');
    }
  }
}

// Partial baseline rerun. `componentIds` empty = full rerun (same outcome as
// runBaseline but without the pipeline side effects like tab switching).
// `orchestrated=true` means runChat already holds setLoading(true) for the
// whole pending_actions sequence — do not toggle it per call, otherwise the
// user can fire a second chat message in the gap between two reruns.
async function runBaselineForComponents(componentIds, reason, orchestrated) {
  if (!state.project) {
    addMsg('Inget projekt att räkna baslinje på.', 'system');
    return;
  }
  const isFull = !componentIds || componentIds.length === 0;
  const scope = isFull
    ? 'hela baslinjen'
    : 'komponent(er) ' + componentIds.join(', ');
  const reasonNote = reason ? ' (' + reason + ')' : '';
  addMsg('AIda räknar om ' + scope + reasonNote + '...', 'system');
  if (!orchestrated) setLoading(true);
  try {
    const body = {project: state.project};
    if (!isFull) body.component_ids = componentIds;
    const r = await authFetch('/api/baseline', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify(body),
    });
    const d = await r.json();
    if (d.error) {
      addMsg('Fel vid baslinjebberäkning: ' + d.error, 'system');
      setLoading(false);
      return;
    }
    if (isFull) {
      state.baseline = d;
      state.alternatives = null;
      state.selections = {};
      state.reportMarkdown = null;
    } else {
      mergeBaselineDelta(d, new Set(componentIds));
      invalidateDownstreamFor(new Set(componentIds));
    }
    scheduleAutoSave();
    if (activeTab === 'baslinje') renderBaslinjeContent();
    else if (activeTab === 'alternativ') renderAlternativContent();
    addMsg('Baslinje uppdaterad' + (isFull ? '' : ' för ' + componentIds.join(', ')) + '.', 'system');
  } catch (e) {
    addMsg('Fel vid baslinjebberäkning: ' + e.message, 'system');
  } finally {
    if (!orchestrated) setLoading(false);
  }
}

async function runAlternativesForComponents(componentIds, userFeedback, reason, orchestrated) {
  if (!state.project || !state.baseline) {
    addMsg('Saknar projekt eller baslinje för alternativ.', 'system');
    return;
  }
  const isFull = !componentIds || componentIds.length === 0;
  const scope = isFull
    ? 'alla alternativ'
    : 'alternativ för komponent(er) ' + componentIds.join(', ');
  const reasonNote = reason ? ' (' + reason + ')' : '';
  const feedbackNote = userFeedback ? ' Önskemål: ' + userFeedback + '.' : '';
  addMsg('AIda kör om ' + scope + reasonNote + feedbackNote + '...', 'system');
  if (!orchestrated) setLoading(true);
  try {
    const body = {project: state.project, baseline: state.baseline};
    if (!isFull) body.component_ids = componentIds;
    if (userFeedback) body.user_feedback = userFeedback;
    const r = await authFetch('/api/alternatives', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify(body),
    });
    const d = await r.json();
    if (d.error) {
      addMsg('Fel vid alternativsökning: ' + d.error, 'system');
      setLoading(false);
      return;
    }
    if (isFull) {
      state.alternatives = d;
      state.selections = {};
      state.reportMarkdown = null;
    } else {
      mergeAlternativesDelta(d, new Set(componentIds));
      // Invalidate selections for components whose alternatives list just changed.
      const cidSet = new Set(componentIds);
      if (state.selections) {
        Object.keys(state.selections).forEach(cid => { if (cidSet.has(cid)) delete state.selections[cid]; });
      }
      state.reportMarkdown = null;
    }
    scheduleAutoSave();
    if (activeTab === 'alternativ') renderAlternativContent();
    addMsg('Alternativ uppdaterade' + (isFull ? '' : ' för ' + componentIds.join(', ')) + '.', 'system');
  } catch (e) {
    addMsg('Fel vid alternativsökning: ' + e.message, 'system');
  } finally {
    if (!orchestrated) setLoading(false);
  }
}

// Replace baseline entries in place, filtered to the components actually
// requested. cidSet acts as a defensive whitelist: a server returning more
// components than requested (regression or future code path) cannot silently
// overwrite unrelated state.
function mergeBaselineDelta(delta, cidSet) {
  if (!delta || !Array.isArray(delta.components)) return;
  const allowed = (delta.components || []).filter(c => cidSet.has(c.component_id));
  if (!state.baseline || !Array.isArray(state.baseline.components)) {
    state.baseline = {components: allowed};
    return;
  }
  const newById = new Map(allowed.map(c => [c.component_id, c]));
  state.baseline.components = state.baseline.components.map(c => newById.get(c.component_id) || c);
  for (const c of allowed) {
    if (!state.baseline.components.find(x => x.component_id === c.component_id)) {
      state.baseline.components.push(c);
    }
  }
}

function mergeAlternativesDelta(delta, cidSet) {
  if (!delta || !Array.isArray(delta.components)) return;
  const allowed = (delta.components || []).filter(c => cidSet.has(c.component_id));
  if (!state.alternatives || !Array.isArray(state.alternatives.components)) {
    state.alternatives = {components: allowed, commentary: delta.commentary || ''};
    return;
  }
  const newById = new Map(allowed.map(c => [c.component_id, c]));
  state.alternatives.components = state.alternatives.components.map(c => newById.get(c.component_id) || c);
  for (const c of allowed) {
    if (!state.alternatives.components.find(x => x.component_id === c.component_id)) {
      state.alternatives.components.push(c);
    }
  }
  if (delta.commentary) state.alternatives.commentary = delta.commentary;
}

// When baseline is rerun for a subset of components, downstream (alternatives,
// selections, report) for those same components is stale by definition.
function invalidateDownstreamFor(cidSet) {
  if (state.alternatives && Array.isArray(state.alternatives.components)) {
    state.alternatives.components = state.alternatives.components.filter(c => !cidSet.has(c.component_id));
    if (state.alternatives.components.length === 0) state.alternatives = null;
  }
  if (state.selections) {
    Object.keys(state.selections).forEach(cid => { if (cidSet.has(cid)) delete state.selections[cid]; });
  }
  state.reportMarkdown = null;
}

// Apply state deltas returned by the chat agent.
function applyAgentStateUpdates(updates) {
  if (!updates || typeof updates !== 'object') return;
  let touched = false;
  // If the agent returns a fresh baseline/alternatives bag, skaling was applied
  // and no stale warning is needed. Otherwise we flag stale on project mutation.
  let baselineStale = false;
  let altsStale = false;
  let materialChanged = false;

  if (updates.project) {
    const prevIds = new Set((state.project?.components || []).map(c => c.id));
    const newIds = new Set((updates.project.components || []).map(c => c.id));
    const sameIds = prevIds.size === newIds.size && [...prevIds].every(id => newIds.has(id));
    // Detect material/category change vs pure quantity change. If category or
    // name changed on any matching component, the baseline value for it is
    // stale and only a recompute restores correctness.
    if (state.project && sameIds) {
      const prevById = new Map(state.project.components.map(c => [c.id, c]));
      for (const c of updates.project.components) {
        const p = prevById.get(c.id);
        if (p && (p.category !== c.category || p.name !== c.name)) {
          materialChanged = true;
          break;
        }
      }
    } else if (state.project && !sameIds) {
      // Added or removed component => baseline coverage changed.
      materialChanged = true;
    }
    state.project = updates.project;
    touched = true;
    if (!('baseline' in updates)) baselineStale = true;
    if (!('alternatives' in updates)) altsStale = true;
    // If components were removed, mirror in baseline/alternatives/selections.
    if (!sameIds) {
      if (state.baseline && state.baseline.components) {
        state.baseline.components = state.baseline.components.filter(c => newIds.has(c.component_id));
      }
      if (state.alternatives && state.alternatives.components) {
        state.alternatives.components = state.alternatives.components.filter(c => newIds.has(c.component_id));
      }
      if (state.selections) {
        Object.keys(state.selections).forEach(cid => { if (!newIds.has(cid)) delete state.selections[cid]; });
      }
    }
  }

  if (updates.baseline) {
    state.baseline = updates.baseline;
    touched = true;
  }
  if (updates.alternatives) {
    state.alternatives = updates.alternatives;
    touched = true;
  }
  if (updates.selections) {
    state.selections = updates.selections;
    touched = true;
  }

  if (touched) {
    // Re-render active tab to reflect changes.
    if (activeTab === 'projekt' && state.project) renderProjektContent();
    else if (activeTab === 'baslinje' && state.baseline) renderBaslinjeContent();
    else if (activeTab === 'alternativ' && state.alternatives) renderAlternativContent();
    scheduleAutoSave();
  }

  // Surface staleness only when the agent couldn't scale — i.e. project changed but baseline/alternatives didn't come back.
  // For material/category changes we surface a concrete action button so the user does not have to guess
  // (chat agent's prompt promises a "Räkna om baslinjen"-button after such changes).
  if (baselineStale && state.baseline) {
    if (materialChanged) {
      renderRecomputeBaselineAction();
    } else {
      addMsg('⚠️ Baslinjen är nu inaktuell efter ändringen. Kör om den för att få nya värden.', 'system');
    }
  }
  if (altsStale && state.alternatives) {
    if (materialChanged) {
      renderRecomputeAlternativesAction();
    } else {
      addMsg('⚠️ Alternativen är nu inaktuella efter ändringen. Kör om dem för aktuella förslag.', 'system');
    }
  }
}

// Action-button row inside the chat. Used after material/category change so the user can recompute
// without typing or hunting for a hidden control.
function renderActionRow(label, handler) {
  document.querySelectorAll('.msg.action-row[data-label="' + label + '"]').forEach(el => el.remove());
  const el = document.createElement('div');
  el.className = 'msg bot action-row';
  el.dataset.label = label;
  const btn = document.createElement('button');
  btn.type = 'button';
  btn.className = 'action-btn';
  btn.textContent = label;
  btn.onclick = async () => {
    btn.disabled = true;
    try { await handler(); } finally { el.remove(); }
  };
  el.appendChild(btn);
  document.getElementById('messages').appendChild(el);
  el.scrollIntoView({behavior:'smooth'});
}

function renderRecomputeBaselineAction() {
  renderActionRow('Räkna om baslinjen', async () => { await runBaseline(); });
}

function renderRecomputeAlternativesAction() {
  renderActionRow('Sök nya alternativ', async () => { await runAlternatives(); });
}

// === Helpers ===
// About modal (Feature 5)
function openAbout() { document.getElementById('aboutModal').style.display = 'flex'; }
function closeAbout() { document.getElementById('aboutModal').style.display = 'none'; }
document.addEventListener('keydown', e => { if (e.key === 'Escape') closeAbout(); });

// Reasoning toggle (Feature 2)
function toggleReasoning(id, e) {
  e.stopPropagation();
  const row = document.getElementById('reasoning-' + id);
  if (!row) return;
  const isHidden = row.style.display === 'none';
  row.style.display = isHidden ? '' : 'none';
  e.target.textContent = isHidden ? 'D\u00f6lj' : 'Visa mer';
}

function formatSource(source) {
  if (!source) return '';
  if (source.startsWith('[EPD]')) return '<span class="source-badge source-verified">EPD</span>' + esc(source.replace('[EPD] ', ''));
  if (source.startsWith('[Palats]')) return '<span class="source-badge source-verified">Palats</span>' + esc(source.replace('[Palats] ', ''));
  if (source.includes('Boverket')) return '<span class="source-badge source-verified">BVK</span>' + esc(source);
  // EPD-typvärde: median of upper-half EPDs by GWP per category — better
  // than Uppskattning, less precise than a single verified EPD. Approximates
  // NollCO2 'Typical' for categories Boverket lacks.
  if (source.includes('EPD-typvärde') || source.includes('EPD-medel') || source.includes('EPD-median')) return '<span class="source-badge source-aggregate">EPD-typvärde</span>' + esc(source);
  if (source.includes('EPD') || source.includes('Environdec')) return '<span class="source-badge source-verified">EPD</span>' + esc(source);
  if (source.startsWith('[Uppskattning]')) return '<span class="source-badge source-estimate">Est.</span>' + esc(source.replace('[Uppskattning] ', ''));
  if (source.includes('Uppskattning')) return '<span class="source-badge source-estimate">Est.</span>' + esc(source);
  return esc(source);
}

function getTypeBadge(alt) {
  if (alt.alternative_type === 'info') return '<span class="type-badge" style="background:var(--kk-gray-100);color:var(--kk-gray-500)">Info</span>';
  if (alt.alternative_type === 'reuse') return '<span class="type-badge type-reuse">\u00c5terbruk</span>';
  if (alt.alternative_type === 'climate_optimized') return '<span class="type-badge type-optimized">Klimatopt.</span>';
  return '<span class="type-badge type-baseline">Baslinje</span>';
}

// === Tab renderers ===
function quantitySourceBadge(src) {
  if (src === 'user_specified') return '<span class="source-badge source-verified" title="Antalet kommer fr\u00e5n din projektbeskrivning">Du angav</span>';
  return '<span class="source-badge source-estimate" title="AIda uppskattade antalet utifr\u00e5n area och byggnadstyp \u2014 granska om n\u00e5got verkar fel">AIda uppskattat</span>';
}

// Snapshot of the inferred text at the moment edit was opened. Restores
// on Cancel even if state.project was swapped by a concurrent chat-agent
// update mid-edit.
let _naEditSnapshot = null;

function renderNeedsAnalysis(na) {
  // Editorial pairing: user voice (gray) \u2194 AIda voice (red), with vertical
  // accent stripes and a clear visual transition between them.
  const hasAny = na && (
    (na.from_user || '').trim() ||
    (na.inferred || '').trim() ||
    (Array.isArray(na.assumptions) && na.assumptions.length) ||
    (Array.isArray(na.would_clarify) && na.would_clarify.length)
  );
  if (!hasAny) {
    return '<div class="needs-card">' +
      '<div class="needs-card-head"><div class="needs-card-title">AIdas behovsanalys</div></div>' +
      '<div class="needs-empty">Ingen behovsanalys finns f\u00f6r det h\u00e4r projektet \u2014 analysen tillkom efter att projektet skapades. K\u00f6ra om intake i chatten ger en analys som styr alternativvalet.</div></div>';
  }
  const fromUser = na.from_user || '';
  const inferred = na.inferred || '';
  const assumptions = Array.isArray(na.assumptions) ? na.assumptions : [];
  const clarify = Array.isArray(na.would_clarify) ? na.would_clarify : [];

  let html = '<div class="needs-card">';
  html += '<div class="needs-card-head">';
  html += '<div class="needs-card-title">AIdas behovsanalys</div>';
  html += '<div class="needs-card-sub">Granska f\u00f6re baslinje \u00b7 korrigera om AIdas l\u00e4sning \u00e4r fel</div>';
  html += '</div>';
  html += '<div class="needs-body">';

  // User voice block
  html += '<div class="voice-block voice-user">';
  html += '<div class="voice-label"><span class="dot"></span>Du sa</div>';
  html += '<div class="voice-text">' + (fromUser ? esc(fromUser) : '<em class="empty">(inget direkt fr\u00e5n din beskrivning)</em>') + '</div>';
  html += '</div>';

  // Transition
  html += '<div class="voice-transition">\u2193 AIdas l\u00e4sning av detta</div>';

  // AIda voice block with edit affordance
  html += '<div class="voice-block voice-aida" id="aidaVoiceBlock">';
  html += '<div class="voice-aida-actions">';
  html += '<button type="button" class="voice-aida-edit" id="naEditBtn" onclick="toggleNeedsEdit()">';
  html += '<svg viewBox="0 0 16 16" fill="none" stroke="currentColor" stroke-width="1.8"><path d="M11 2l3 3-9 9H2v-3l9-9z"/></svg>';
  html += 'Justera AIdas l\u00e4sning</button>';
  html += '</div>';
  html += '<div class="voice-label"><span class="dot"></span>AIda tolkar</div>';
  html += '<div class="voice-text" id="naInferredView">' + (inferred ? esc(inferred) : '<em class="empty">(ingen inferens)</em>') + '</div>';
  // Empty textarea \u2014 value populated via .value to preserve quotes/ampersands
  html += '<textarea class="voice-aida-textarea" id="naInferredEdit"></textarea>';
  html += '<div class="voice-aida-edit-actions" id="naEditActions">';
  html += '<button type="button" class="btn-na-cancel" onclick="cancelNeedsEdit()">Avbryt</button>';
  html += '<button type="button" class="btn-na-save" onclick="saveNeedsEdit()">Spara</button>';
  html += '</div>';
  html += '</div>';

  // Meta row \u2014 assumptions + would_clarify
  if (assumptions.length || clarify.length) {
    html += '<div class="needs-meta-row">';
    html += '<div>';
    if (assumptions.length) {
      html += '<div class="needs-meta-label">Antaganden AIda gjort</div>';
      html += '<ul class="needs-meta-list needs-meta-assumptions">';
      assumptions.forEach(a => { html += '<li>' + esc(a) + '</li>'; });
      html += '</ul>';
    }
    html += '</div><div>';
    if (clarify.length) {
      html += '<div class="needs-meta-label">AIda hade g\u00e4rna vetat</div>';
      html += '<ul class="needs-meta-list needs-meta-clarify">';
      clarify.forEach(q => { html += '<li>' + esc(q) + '</li>'; });
      html += '</ul>';
    }
    html += '</div></div>';
  }

  html += '</div></div>';
  return html;
}

function _populateNeedsTextarea() {
  // Populate the textarea via .value (not innerHTML) so the raw inferred
  // text \u2014 including quotes and ampersands \u2014 is preserved verbatim.
  const ta = document.getElementById('naInferredEdit');
  if (!ta) return;
  const na = state.project && state.project.needs_analysis;
  ta.value = (na && na.inferred) || '';
}

function toggleNeedsEdit() {
  const ta = document.getElementById('naInferredEdit');
  const block = document.getElementById('aidaVoiceBlock');
  if (!ta || !block) return;
  _naEditSnapshot = ta.value;  // remember pre-edit value
  block.classList.add('is-editing');
  ta.focus();
}

function cancelNeedsEdit() {
  const ta = document.getElementById('naInferredEdit');
  const block = document.getElementById('aidaVoiceBlock');
  if (!ta || !block) return;
  // Restore from the pre-edit snapshot \u2014 not from state.project, which may
  // have been mutated by a concurrent chat-agent update during the edit.
  ta.value = _naEditSnapshot != null ? _naEditSnapshot : ta.value;
  _naEditSnapshot = null;
  block.classList.remove('is-editing');
}

function saveNeedsEdit() {
  const ta = document.getElementById('naInferredEdit');
  if (!ta) return;
  const newVal = ta.value.trim();
  if (!state.project.needs_analysis) state.project.needs_analysis = {from_user:'',inferred:'',assumptions:[],would_clarify:[]};
  state.project.needs_analysis.inferred = newVal;
  _naEditSnapshot = null;
  if (typeof scheduleAutoSave === 'function') scheduleAutoSave();
  renderProjektContent();
}

function renderProjektContent() {
  const d = state.project;
  let html = '<div class="section-title">Projektinformation</div>';
  html += renderNeedsAnalysis(d.needs_analysis);
  html += '<div class="comp-card"><div class="comp-card-header"><h3>' + esc(d.building_type) + ', ' + esc(d.area_bta) + ' m\u00b2 BTA' + (d.name ? ' (' + esc(d.name) + ')' : '') + '</h3></div>';
  html += '<table class="comp-table"><thead><tr><th>Komponent</th><th>Antal</th><th>Enhet</th><th>Kategori</th><th>K\u00e4lla</th></tr></thead><tbody>';
  d.components.forEach(c => {
    const nameCell = '<div style="font-weight:500">' + esc(c.name) + '</div>' +
      (c.usage_context ? '<div class="usage-context"><span class="usage-context-label">Anv\u00e4ndning</span>' + esc(c.usage_context) + '</div>' : '');
    html += '<tr>' +
      '<td>' + nameCell + '</td>' +
      '<td>' + esc(c.quantity) + '</td>' +
      '<td>' + esc(c.unit) + '</td>' +
      '<td>' + esc(c.category || '\u2013') + '</td>' +
      '<td>' + quantitySourceBadge(c.quantity_source) + '</td>' +
      '</tr>';
  });
  html += '</tbody></table></div>';
  if (d.description) {
    html += '<div class="comp-card" style="margin-top:12px"><div class="comp-card-header"><h3>Beskrivning</h3></div><div style="padding:12px 16px;font-size:13px;color:var(--kk-gray-500);line-height:1.5">' + esc(d.description) + '</div></div>';
  }
  document.getElementById('resultContent').innerHTML = html;
  _populateNeedsTextarea();
}

function renderBaslinjeContent() {
  const d = state.baseline;
  const total = d.components.reduce((s,c) => s + c.co2e_kg, 0);
  const totalCost = d.components.reduce((s,c) => s + c.cost_sek, 0);
  let html = '<div class="section-title">Baslinje (NollCO2-metoden)</div>';
  html += '<div class="method-label">Klimatmetod: GWP-fossil, livscykelskedena A1-A3 (Boverkets klimatdatabas)</div>';
  html += '<div class="source-legend"><span><span class="source-badge source-verified">EPD</span> Verifierad k\u00e4lla</span><span><span class="source-badge source-aggregate">EPD-typvärde</span> Kategori-typvärde (övre halvan)</span><span><span class="source-badge source-estimate">Est.</span> Uppskattning</span></div>';
  html += '<div class="summary">';
  html += '<div class="card"><div class="card-title">Total CO\u2082e</div><div class="value">' + Math.round(total).toLocaleString('sv') + '</div><div class="sublabel">kg CO\u2082e</div></div>';
  html += '<div class="card"><div class="card-title">Total kostnad</div><div class="value">' + Math.round(totalCost).toLocaleString('sv') + '</div><div class="sublabel">SEK</div></div>';
  html += '<div class="card"><div class="card-title">Komponenter</div><div class="value">' + d.components.length + '</div><div class="sublabel">st</div></div>';
  html += '</div>';
  html += '<div class="comp-card"><div class="comp-card-header"><h3>Per komponent</h3></div>';
  html += '<table class="comp-table"><thead><tr><th>Komponent</th><th style="text-align:right">CO\u2082e (kg)</th><th>Klimatk\u00e4lla</th><th style="text-align:right">Kostnad (SEK)</th><th>Prisk\u00e4lla</th></tr></thead><tbody>';
  d.components.forEach(c => {
    // Frame Boverket products as proxies — Boverket has ~200 material entries,
    // not building-component entries, so almost all matches are material-based
    // proxies (e.g. vinylgolv → "Takduk, PVC"). Showing the bare product name
    // reads as a category mismatch.
    const productLine = c.boverket_product ? '<div style="font-size:11px;color:var(--kk-gray-500);margin-top:3px;font-style:italic"><span style="font-style:normal;font-weight:500;color:var(--kk-gray-400);font-size:9.5px;letter-spacing:0.8px;text-transform:uppercase;display:block;margin-bottom:1px">Materialproxy</span>' + esc(c.boverket_product) + '</div>' : '';
    html += '<tr><td style="font-weight:500">' + esc(c.component_name) + '</td><td style="text-align:right">' + Math.round(c.co2e_kg).toLocaleString('sv') + '</td><td style="font-size:11px">' + formatSource(c.source) + productLine + '</td><td style="text-align:right">' + Math.round(c.cost_sek).toLocaleString('sv') + '</td><td style="font-size:11px">' + esc(c.cost_source || '') + '</td></tr>';
  });
  html += '</tbody></table></div>';
  document.getElementById('resultContent').innerHTML = html;
}

function renderAlternativContent() {
  const data = state.alternatives;
  let html = '<div class="section-title">J\u00e4mf\u00f6relse per komponent</div>';
  html += '<div class="method-label">Klimatmetod: GWP-fossil, livscykelskedena A1-A3 (Boverkets klimatdatabas)</div>';
  html += '<div class="source-legend"><span><span class="source-badge source-verified">EPD</span> Verifierad k\u00e4lla</span><span><span class="source-badge source-aggregate">EPD-typvärde</span> Kategori-typvärde (övre halvan)</span><span><span class="source-badge source-estimate">Est.</span> Uppskattning</span></div>';
  const projComps = (state.project && state.project.components) || [];
  data.components.forEach(comp => {
    const pc = projComps.find(p => p.id === comp.component_id);
    const qtyLabel = pc ? esc(pc.quantity) + ' ' + esc(pc.unit) + ' ' + quantitySourceBadge(pc.quantity_source) : '';
    const usageBlock = (pc && pc.usage_context) ? '<div class="usage-context"><span class="usage-context-label">Användning</span>' + esc(pc.usage_context) + '</div>' : '';
    const header = '<h3>' + esc(comp.component_name) + '</h3>' + (qtyLabel ? '<div style="font-size:12px;color:var(--kk-gray-500);margin-top:2px">Antal: ' + qtyLabel + '</div>' : '') + usageBlock;
    html += '<div class="comp-card"><div class="comp-card-header">' + header + '</div>';
    html += '<table class="comp-table"><thead><tr><th style="width:32px"></th><th>Typ</th><th>Material</th><th>K\u00e4lla</th><th style="text-align:right">CO\u2082e (kg)</th><th style="text-align:right">Kostnad</th><th></th></tr></thead><tbody>';
    const blSel = state.selections[comp.component_id] && state.selections[comp.component_id].selected_alternative.name === 'Baslinje';
    // Look up the Boverket product used for this component's baseline so the
    // baseline row in the alternatives table shows the actual proxy product
    // (e.g. "Takduk, PVC") rather than just the generic "Konventionellt".
    const blBaselineComp = (state.baseline && state.baseline.components) ? state.baseline.components.find(b => b.component_id === comp.component_id) : null;
    const blProduct = (blBaselineComp && blBaselineComp.boverket_product) ? blBaselineComp.boverket_product : '';
    const blSource = (blBaselineComp && blBaselineComp.source) ? blBaselineComp.source : 'NollCO2';
    const blMaterialCell = blProduct
      ? '<div style="font-weight:500">Konventionellt</div><div style="font-size:11px;color:var(--kk-gray-500);font-style:italic;margin-top:2px"><span style="font-style:normal;font-weight:500;color:var(--kk-gray-400);font-size:9.5px;letter-spacing:0.8px;text-transform:uppercase;display:block;margin-bottom:1px">Materialproxy</span>' + esc(blProduct) + '</div>'
      : '<div style="font-weight:500">Konventionellt</div>';
    html += '<tr class="alt-row' + (blSel ? ' selected' : '') + '" data-comp="' + comp.component_id + '" data-alt="baseline">' +
      '<td><input type="radio" name="' + comp.component_id + '"' + (blSel ? ' checked' : '') + '></td>' +
      '<td><span class="type-badge type-baseline">Baslinje</span></td>' +
      '<td>' + blMaterialCell + '</td><td style="font-size:11px">' + (blSource.includes('Boverket') ? '<span class="source-badge source-verified">BVK</span>' : '<span class="source-badge source-estimate">Est.</span>') + ' NollCO2</td>' +
      '<td style="text-align:right">' + Math.round(comp.baseline_co2e_kg) + '</td>' +
      '<td style="text-align:right">' + Math.round(comp.baseline_cost_sek).toLocaleString('sv') + ' kr</td><td></td></tr>';
    comp.alternatives.forEach((alt, i) => {
      const rowId = comp.component_id + '_' + i;
      if (alt.alternative_type === 'info') {
        html += '<tr style="opacity:0.6">' +
          '<td></td>' +
          '<td>' + getTypeBadge(alt) + '</td>' +
          '<td colspan="4" style="font-size:12px;color:var(--kk-gray-500)">' + esc(alt.name) + '</td>' +
          '<td>' + (alt.reasoning ? '<button class="reasoning-toggle" onclick="toggleReasoning(\'' + rowId + '\',event)">Visa mer</button>' : '') + '</td></tr>';
        if (alt.reasoning) {
          html += '<tr class="reasoning-row" id="reasoning-' + rowId + '" style="display:none"><td colspan="7">' + esc(alt.reasoning) + '</td></tr>';
        }
        return;
      }
      const saving = comp.baseline_co2e_kg > 0 ? Math.round((1 - alt.co2e_kg / comp.baseline_co2e_kg) * 100) : 0;
      const isSel = state.selections[comp.component_id] && state.selections[comp.component_id].selected_alternative.name === alt.name;
      // Decompose total for reuse alternatives where units match (no trailing *).
      // Lets the user see e.g. "45 st \u00d7 320 kr = 14 400 kr" inline rather than
      // hidden in Visa mer \u2014 answers Johanna's "varf\u00f6r 45 lampor" without exposing
      // adjustability yet.
      const showBreakdown = alt.alternative_type === 'reuse' && !alt.name.endsWith('*') && pc && pc.quantity > 0 && alt.cost_sek > 0;
      const perUnit = showBreakdown ? Math.round(alt.cost_sek / pc.quantity) : 0;
      const costCell = alt.name.endsWith('*')
        ? Math.round(alt.cost_sek).toLocaleString('sv') + ' kr/st *'
        : (showBreakdown
          ? '<div style="line-height:1.3">' + Math.round(alt.cost_sek).toLocaleString('sv') + ' kr<div style="font-size:10px;color:var(--kk-gray-500)">' + esc(String(pc.quantity)) + ' \u00d7 ' + perUnit.toLocaleString('sv') + ' kr</div></div>'
          : Math.round(alt.cost_sek).toLocaleString('sv') + ' kr');
      html += '<tr class="alt-row' + (isSel ? ' selected' : '') + '" data-comp="' + comp.component_id + '" data-alt="' + i + '">' +
        '<td><input type="radio" name="' + comp.component_id + '"' + (isSel ? ' checked' : '') + '></td>' +
        '<td>' + getTypeBadge(alt) + '</td>' +
        '<td style="font-weight:500">' + esc(alt.name) + '</td>' +
        '<td style="font-size:11px">' + formatSource(alt.source) + '</td>' +
        '<td style="text-align:right">' + Math.round(alt.co2e_kg) + ' <span style="color:' + (saving >= 0 ? 'var(--green-saving)' : 'var(--kk-red-orange)') + ';font-size:11px">' + (saving >= 0 ? '\u2193' : '\u2191') + Math.abs(saving) + '%</span></td>' +
        '<td style="text-align:right">' + costCell + '</td>' +
        '<td>' + (alt.reasoning ? '<button class="reasoning-toggle" onclick="toggleReasoning(\'' + rowId + '\',event)">Visa mer</button>' : '') + '</td></tr>';
      if (alt.reasoning) {
        html += '<tr class="reasoning-row" id="reasoning-' + rowId + '" style="display:none"><td colspan="7">' + esc(alt.reasoning) + '</td></tr>';
      }
    });
    html += '</tbody></table></div>';
  });
  // Check if any alternatives have per-article pricing
  const hasPerArticle = data.components.some(c => c.alternatives.some(a => a.name.endsWith('*')));
  if (hasPerArticle) {
    html += '<div style="font-size:12px;color:var(--kk-gray-500);margin:8px 0;font-style:italic">* Pris per artikel (yta per artikel ok\u00e4nd). Se \u201cVisa mer\u201d f\u00f6r detaljer.</div>';
  }
  html += '<div id="summaryArea"></div>';
  html += '<button class="btn" id="reportBtn" onclick="generateReport()" disabled title="V\u00e4lj ett alternativ per komponent">Generera rapport</button>';
  html += '<div id="missingHint" style="font-size:12px;color:var(--kk-gray-500);margin-top:6px;font-style:italic"></div>';
  document.getElementById('resultContent').innerHTML = html;
  // Bind click handlers
  document.querySelectorAll('.alt-row').forEach(row => {
    row.onclick = function() { selectAlt(this.dataset.comp, this.dataset.alt, this); };
  });
  if (Object.keys(state.selections).length > 0) updateSummary();
}

function renderRapportContent() {
  let html = '<div class="report-area">' + renderMd(state.reportMarkdown) + '</div>';
  html += '<div style="margin-top:12px;display:flex;gap:8px">';
  html += '<button class="btn" id="dlDocxBtn"><svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" style="vertical-align:-2px;margin-right:4px"><path d="M21 15v4a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2v-4"/><polyline points="7 10 12 15 17 10"/><line x1="12" y1="15" x2="12" y2="3"/></svg>Ladda ner Word (.docx)</button>';
  html += '<button class="btn btn-secondary" id="dlBtn"><svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" style="vertical-align:-2px;margin-right:4px"><path d="M21 15v4a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2v-4"/><polyline points="7 10 12 15 17 10"/><line x1="12" y1="15" x2="12" y2="3"/></svg>Ladda ner (.md)</button>';
  html += '</div>';
  document.getElementById('resultContent').innerHTML = html;
  document.getElementById('dlBtn').onclick = () => {
    const blob = new Blob([state.reportMarkdown], {type:'text/markdown'});
    const a = document.createElement('a'); a.href = URL.createObjectURL(blob); a.download = 'aida-rapport.md'; a.click();
  };
  document.getElementById('dlDocxBtn').onclick = async () => {
    const btn = document.getElementById('dlDocxBtn');
    btn.disabled = true; btn.textContent = 'Skapar dokument...';
    try {
      const r = await authFetch('/api/report/docx', {method:'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify({markdown: state.reportMarkdown})});
      if (!r.ok) { const d = await r.json(); alert('Fel: ' + d.error); return; }
      const blob = await r.blob();
      const today = new Date().toISOString().slice(0,10);
      const a = document.createElement('a'); a.href = URL.createObjectURL(blob); a.download = 'AIda_rapport_' + today + '.docx'; a.click();
    } catch(e) { alert('Fel: ' + e.message); }
    finally { btn.disabled = false; btn.innerHTML = '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" style="vertical-align:-2px;margin-right:4px"><path d="M21 15v4a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2v-4"/><polyline points="7 10 12 15 17 10"/><line x1="12" y1="15" x2="12" y2="3"/></svg>Ladda ner Word (.docx)'; }
  };
}

// === Selection handling ===
function selectAlt(compId, altIdx, row) {
  row.closest('table').querySelectorAll('.alt-row').forEach(r => r.classList.remove('selected'));
  row.classList.add('selected');
  row.querySelector('input[type=radio]').checked = true;
  const comp = state.alternatives.components.find(c => c.component_id === compId);
  if (altIdx === 'baseline') {
    state.selections[compId] = { id: compId, name: comp.component_name,
      selected_alternative: {name:'Baslinje', co2e_kg: comp.baseline_co2e_kg, cost_sek: comp.baseline_cost_sek, source:'NollCO2'},
      baseline_co2e_kg: comp.baseline_co2e_kg, baseline_cost_sek: comp.baseline_cost_sek };
  } else {
    const alt = comp.alternatives[parseInt(altIdx)];
    state.selections[compId] = { id: compId, name: comp.component_name,
      selected_alternative: {name: alt.name, co2e_kg: alt.co2e_kg, cost_sek: alt.cost_sek, source: alt.source},
      baseline_co2e_kg: comp.baseline_co2e_kg, baseline_cost_sek: comp.baseline_cost_sek };
  }
  updateSummary();
  scheduleAutoSave();
}

function updateSummary() {
  const sels = Object.values(state.selections);
  const totalCo2 = sels.reduce((s,c) => s + c.selected_alternative.co2e_kg, 0);
  const totalCost = sels.reduce((s,c) => s + c.selected_alternative.cost_sek, 0);
  const blCo2 = sels.reduce((s,c) => s + c.baseline_co2e_kg, 0);
  const blCost = sels.reduce((s,c) => s + c.baseline_cost_sek, 0);
  const co2Diff = totalCo2 - blCo2;
  const co2Pct = blCo2 > 0 ? Math.round(Math.abs(co2Diff) / blCo2 * 100) : 0;
  const co2Arrow = co2Diff <= 0 ? '\u2193' : '\u2191';
  const costDiff = totalCost - blCost;
  const costPct = blCost > 0 ? Math.round(Math.abs(costDiff) / blCost * 100) : 0;
  const costArrow = costDiff <= 0 ? '\u2193' : '\u2191';
  document.getElementById('summaryArea').innerHTML =
    '<div class="summary">' +
    '<div class="card' + (co2Diff <= 0 ? ' saving' : '') + '"><div class="card-title">Klimatp\u00e5verkan</div><div class="value">' + Math.round(totalCo2).toLocaleString('sv') + '</div><div class="sublabel">kg CO\u2082e (' + co2Arrow + co2Pct + '% vs baslinje)</div></div>' +
    '<div class="card' + (costDiff <= 0 ? ' saving' : '') + '"><div class="card-title">Kostnad</div><div class="value">' + Math.round(totalCost).toLocaleString('sv') + '</div><div class="sublabel">SEK (' + costArrow + costPct + '% vs baslinje)</div></div>' +
    '<div class="card"><div class="card-title">Baslinje</div><div class="value">' + Math.round(blCo2).toLocaleString('sv') + '</div><div class="sublabel">kg CO\u2082e | ' + Math.round(blCost).toLocaleString('sv') + ' SEK</div></div>' +
    '</div>';
  const allSelected = state.alternatives.components.every(c => state.selections[c.component_id]);
  document.getElementById('reportBtn').disabled = !allSelected;
  const hintEl = document.getElementById('missingHint');
  if (hintEl) {
    if (!allSelected) {
      const missing = state.alternatives.components.filter(c => !state.selections[c.component_id]).map(c => c.component_name);
      hintEl.textContent = 'V\u00e4lj alternativ f\u00f6r: ' + missing.join(', ');
    } else { hintEl.textContent = ''; }
  }
}

// === Supabase auth + persistence ===
const HAS_SUPABASE = {{ 'true' if has_supabase else 'false' }};
const SUPABASE_URL = {{ supabase_url|tojson }};
const SUPABASE_ANON_KEY = {{ supabase_anon_key|tojson }};
let supabaseClient = null;
let currentUser = null;
let currentAnalysisId = null;
let isSignup = false;
let saveTimeout = null;
let saveInProgress = false;

// Auth-aware fetch wrapper with safe JSON parsing
async function authFetch(url, options) {
  options = options || {};
  options.headers = options.headers || {};
  if (supabaseClient) {
    const sess = await supabaseClient.auth.getSession();
    if (sess.data.session) {
      options.headers['Authorization'] = 'Bearer ' + sess.data.session.access_token;
    }
  }
  const resp = await fetch(url, options);
  // Wrap .json() to catch non-JSON responses (e.g. Vercel timeout HTML pages)
  const origJson = resp.json.bind(resp);
  resp.json = async () => {
    const text = await resp.clone().text();
    try { return JSON.parse(text); }
    catch(e) {
      if (resp.status === 504 || resp.status === 502 || text.includes('FUNCTION_INVOCATION_TIMEOUT'))
        return {error: 'Analysen tog f\u00f6r l\u00e5ng tid. F\u00f6rs\u00f6k igen, eller f\u00f6renkla projektbeskrivningen.'};
      if (resp.status >= 500)
        return {error: 'Serverfel (' + resp.status + '). F\u00f6rs\u00f6k igen om en stund.'};
      return {error: 'Ov\u00e4ntat svar fr\u00e5n servern. F\u00f6rs\u00f6k igen.'};
    }
  };
  return resp;
}

// No-op when Supabase not configured
function scheduleAutoSave() {
  if (!HAS_SUPABASE || !currentUser) return;
  if (saveTimeout) clearTimeout(saveTimeout);
  saveTimeout = setTimeout(autoSave, 2000);
}

async function autoSave() {
  if (!supabaseClient || !currentUser || saveInProgress) return;
  saveInProgress = true;
  const indicator = document.getElementById('saveIndicator');
  if (indicator) { indicator.textContent = 'Sparar...'; indicator.style.display = 'inline'; indicator.style.color = 'var(--kk-gray-400)'; }
  const analysisData = {
    name: state.project ? (state.project.name || state.project.building_type || 'Nytt projekt') : 'Nytt projekt',
    status: state.step,
    project_data: state.project,
    baseline_data: state.baseline,
    alternatives_data: state.alternatives,
    selections_data: Object.keys(state.selections).length > 0 ? state.selections : null,
    report_markdown: state.reportMarkdown,
  };
  try {
    if (currentAnalysisId) {
      const r = await authFetch('/api/analyses/' + currentAnalysisId, {
        method: 'PUT', headers: {'Content-Type':'application/json'},
        body: JSON.stringify(analysisData),
      });
      await r.json();
    } else {
      const r = await authFetch('/api/analyses', {
        method: 'POST', headers: {'Content-Type':'application/json'},
        body: JSON.stringify(analysisData),
      });
      const result = await r.json();
      if (result && result.id) {
        currentAnalysisId = result.id;
        await loadAnalysesList();
      }
    }
    if (indicator) { indicator.textContent = 'Sparat'; setTimeout(() => { indicator.style.display = 'none'; }, 2000); }
  } catch (e) {
    console.error('Auto-save failed:', e);
    if (indicator) { indicator.textContent = 'Sparfel'; indicator.style.color = 'var(--kk-dark-red)'; }
  }
  finally { saveInProgress = false; }
}

if (HAS_SUPABASE && SUPABASE_URL && SUPABASE_ANON_KEY) {
  supabaseClient = supabase.createClient(SUPABASE_URL, SUPABASE_ANON_KEY);
  initAuth();
}

async function initAuth() {
  const { data: { session } } = await supabaseClient.auth.getSession();
  if (session) { onLogin(session); }
  else { showAuth(); }
  supabaseClient.auth.onAuthStateChange((event, session) => {
    if (event === 'SIGNED_IN' && session) {
      if (!currentUser) onLogin(session);
    }
    else if (event === 'SIGNED_OUT') { currentUser = null; showAuth(); }
  });
}

function showAuth() {
  document.getElementById('authOverlay').style.display = 'flex';
  document.getElementById('appContainer').style.display = 'none';
}

function showApp() {
  document.getElementById('authOverlay').style.display = 'none';
  document.getElementById('appContainer').style.display = '';
}

async function onLogin(session) {
  currentUser = session.user;
  document.getElementById('userEmail').textContent = currentUser.email;
  showApp();
  const list = await loadAnalysesList();
  if (list && list.length > 0) { await loadAnalysis(list[0].id); }
}

async function handleAuth() {
  const email = document.getElementById('authEmail').value.trim();
  const password = document.getElementById('authPassword').value;
  const errorEl = document.getElementById('authError');
  errorEl.style.display = 'none';
  if (!email || !password) {
    errorEl.textContent = 'Fyll i e-post och lösenord';
    errorEl.style.display = 'block';
    return;
  }
  document.getElementById('authSubmitBtn').disabled = true;
  try {
    const result = isSignup
      ? await supabaseClient.auth.signUp({ email, password })
      : await supabaseClient.auth.signInWithPassword({ email, password });
    if (result.error) {
      const AUTH_ERRORS = {'Invalid login credentials':'Fel e-post eller l\u00f6senord.','Email not confirmed':'Bekr\u00e4fta din e-post innan du loggar in.','User already registered':'Det finns redan ett konto med den e-postadressen.','Password should be at least 6 characters':'L\u00f6senordet m\u00e5ste vara minst 6 tecken.'};
      errorEl.textContent = AUTH_ERRORS[result.error.message] || result.error.message;
      errorEl.style.display = 'block';
    } else if (isSignup && !result.data.session) {
      errorEl.textContent = 'Kolla din e-post för bekräftelselänk';
      errorEl.style.display = 'block';
      errorEl.style.color = 'var(--green-saving)';
    }
  } catch (e) {
    errorEl.textContent = e.message;
    errorEl.style.display = 'block';
  }
  document.getElementById('authSubmitBtn').disabled = false;
}

function toggleAuthMode(e) {
  e.preventDefault();
  isSignup = !isSignup;
  document.getElementById('authSubmitBtn').textContent = isSignup ? 'Skapa konto' : 'Logga in';
  document.getElementById('authToggleText').textContent = isSignup ? 'Har redan konto?' : 'Inget konto?';
  document.getElementById('authToggleLink').textContent = isSignup ? 'Logga in' : 'Skapa konto';
  document.getElementById('authError').style.display = 'none';
}

async function handleLogout() {
  await supabaseClient.auth.signOut();
  currentUser = null;
  currentAnalysisId = null;
  showAuth();
}

// === Project dropdown ===
function toggleProjectMenu() {
  const m = document.getElementById('projectMenu');
  const u = document.getElementById('userMenu');
  if (u) u.style.display = 'none';
  m.style.display = m.style.display === 'none' ? 'block' : 'none';
}

// === Project rename ===
function startRenameProject() {
  document.getElementById('projectMenu').style.display = 'none';
  const span = document.getElementById('projectName');
  if (!span) return;
  const current = (state.project && state.project.name) ? state.project.name : span.textContent;
  const input = document.createElement('input');
  input.type = 'text';
  input.className = 'project-rename-input';
  input.value = current === 'Nytt projekt' ? '' : current;
  input.placeholder = 'Projektnamn';
  input.maxLength = 80;
  span.style.display = 'none';
  span.parentNode.insertBefore(input, span);
  input.focus();
  input.select();
  const commit = (save) => {
    if (!input.parentNode) return;
    const next = input.value.trim();
    if (save && next) {
      if (!state.project) state.project = {name: next};
      else state.project.name = next;
      span.textContent = next;
      if (HAS_SUPABASE && currentUser) scheduleAutoSave();
    }
    input.remove();
    span.style.display = '';
  };
  input.addEventListener('keydown', (e) => {
    if (e.key === 'Enter') { e.preventDefault(); commit(true); }
    else if (e.key === 'Escape') { e.preventDefault(); commit(false); }
  });
  input.addEventListener('blur', () => commit(true));
}

function toggleUserMenu() {
  const m = document.getElementById('userMenu');
  const p = document.getElementById('projectMenu');
  if (p) p.style.display = 'none';
  m.style.display = m.style.display === 'none' ? 'block' : 'none';
}

document.addEventListener('click', (e) => {
  const pd = document.getElementById('projectDropdown');
  const ud = document.getElementById('userDropdown');
  if (pd && !e.target.closest('#projectDropdown')) document.getElementById('projectMenu').style.display = 'none';
  if (ud && !e.target.closest('#userDropdown')) document.getElementById('userMenu').style.display = 'none';
});

async function loadAnalysesList() {
  if (!supabaseClient || !currentUser) return null;
  try {
    const r = await authFetch('/api/analyses');
    const list = await r.json();
    const container = document.getElementById('projectList');
    if (!container) return list;
    container.innerHTML = '';
    if (list && list.length > 0) {
      list.forEach(a => {
        const item = document.createElement('div');
        item.className = 'dropdown-item' + (a.id === currentAnalysisId ? ' active' : '');
        item.style.cssText = 'display:flex;align-items:center;justify-content:space-between;cursor:pointer';
        const nameSpan = document.createElement('span');
        nameSpan.textContent = a.name || 'Nytt projekt';
        nameSpan.style.cssText = 'overflow:hidden;text-overflow:ellipsis;white-space:nowrap;flex:1';
        nameSpan.onclick = () => { loadAnalysis(a.id); toggleProjectMenu(); };
        item.appendChild(nameSpan);
        const del = document.createElement('button');
        del.style.cssText = 'background:none;border:none;cursor:pointer;color:var(--kk-gray-400);padding:2px 4px;flex-shrink:0';
        del.title = 'Ta bort';
        del.innerHTML = '\u2715';
        del.onclick = async (e) => {
          e.stopPropagation();
          if (!confirm('Ta bort "' + (a.name || 'Nytt projekt') + '"?')) return;
          try { await authFetch('/api/analyses/' + a.id, {method:'DELETE'}); if (a.id === currentAnalysisId) createNewProject(); await loadAnalysesList(); } catch(ex) { alert('Kunde inte ta bort.'); }
        };
        item.appendChild(del);
        container.appendChild(item);
      });
    } else {
      container.innerHTML = '<div style="padding:8px 16px;font-size:12px;color:var(--kk-gray-400)">Inga projekt ännu</div>';
    }
    return list;
  } catch(e) { console.error('Failed to load list:', e); return null; }
}

async function loadAnalysis(id) {
  if (saveTimeout) { clearTimeout(saveTimeout); saveTimeout = null; }
  try {
    const r = await authFetch('/api/analyses/' + id);
    const data = await r.json();
    if (!data || data.error) return;
    currentAnalysisId = id;
    state.project = data.project_data;
    state.baseline = data.baseline_data;
    state.alternatives = data.alternatives_data;
    state.selections = data.selections_data || {};
    state.reportMarkdown = data.report_markdown;
    state.step = data.status || 'idle';
    document.getElementById('projectName').textContent = data.name || 'Nytt projekt';
    restoreUI();
    await loadAnalysesList();
  } catch(e) { console.error('Failed to load analysis:', e); }
}

function restoreUI() {
  ['projekt','baslinje','alternativ','rapport'].forEach(t => {
    const el = document.getElementById('tab-' + t); if (el) el.disabled = true;
  });
  document.getElementById('progressFill').style.width = '0%';
  document.querySelectorAll('.step-circle').forEach((c, i) => { c.className = 'step-circle'; c.textContent = i + 1; });
  document.querySelectorAll('.step-label').forEach(l => l.className = 'step-label');

  if (state.project) { enableTab('projekt'); setProgressStep('planering'); }
  if (state.baseline) { enableTab('baslinje'); setProgressStep('baslinje'); }
  if (state.alternatives) { enableTab('alternativ'); setProgressStep('sammanstallning'); }
  if (state.reportMarkdown) { enableTab('rapport'); setProgressStep('uppfoljning'); switchTab('rapport'); }
  else if (state.alternatives) { switchTab('alternativ'); }
  else if (state.baseline) { switchTab('baslinje'); }
  else if (state.project) { switchTab('projekt'); }

  const msgs = document.getElementById('messages');
  msgs.innerHTML = '';
  chatLog = [];

  // Try restoring chat from localStorage (survives page reloads)
  let savedChat;
  try { savedChat = JSON.parse(localStorage.getItem(_chatStorageKey())); } catch(e) {}

  if (savedChat && savedChat.length > 0) {
    savedChat.forEach(m => {
      const d = document.createElement('div');
      d.className = 'msg ' + m.cls;
      if (m.cls === 'bot' || m.cls === 'system') { d.innerHTML = renderMd(m.text); }
      else { d.textContent = m.text; }
      // Restore confirm buttons for current step only
      if (m.confirm && (
        (state.step === 'intake_done' && m.confirm.btnLabel.includes('baslinje')) ||
        (state.step === 'baseline_done' && m.confirm.btnLabel.includes('alternativ'))
      )) {
        d.innerHTML += '<div class="confirm-actions"><button class="btn-confirm" onclick="confirmStep()">' + m.confirm.btnLabel + '</button></div><div class="confirm-hint">' + m.confirm.hint + '</div>';
      } else if (m.confirm) {
        // Past confirm — show as completed
        d.innerHTML += '<div class="confirm-actions"><button class="btn-confirm" disabled style="background:var(--kk-gray-200);color:var(--kk-gray-400);cursor:default;pointer-events:none">Bekr\u00e4ftad \u2713</button></div>';
      }
      msgs.appendChild(d);
    });
    chatLog = savedChat;
    const last = msgs.lastElementChild;
    if (last) last.scrollIntoView({behavior:'smooth'});
  } else {
    // Fallback: reconstruct summary from state
    if (!state.project) {
      addMsg('Hej! Beskriv ditt projekt. Ange byggnadstyp, byggnadsår, ungefärlig yta och vilka behoven är.', 'bot');
    } else {
      addMsg('Projekt laddat: ' + (state.project.building_type || 'Okänt') + ', ' + (state.project.area_bta || '?') + ' m\u00b2.', 'bot');
      if (state.step === 'intake_done') {
        const compList = state.project.components.map(c => '- ' + c.name + ' (' + c.quantity + ' ' + c.unit + ')').join('\n');
        addConfirmMsg(
          '**' + state.project.building_type + '**, ' + state.project.area_bta + ' m\u00b2\n\n**Komponenter:**\n' + compList,
          'Bekr\u00e4fta och ber\u00e4kna baslinje \u2192',
          'Skriv i chatten om n\u00e5got inte st\u00e4mmer.'
        );
      } else if (state.step === 'baseline_done') {
        const total = state.baseline.components.reduce((s,c) => s + c.co2e_kg, 0);
        addConfirmMsg(
          'Baslinje klar: **' + Math.round(total).toLocaleString('sv') + ' kg CO\u2082e** totalt f\u00f6r ' + state.baseline.components.length + ' komponenter.',
          'Bekr\u00e4fta och s\u00f6k alternativ \u2192',
          'Skriv i chatten om du vill korrigera n\u00e5got.'
        );
      } else if (state.step === 'alternatives_done') addMsg('V\u00e4lj alternativ per komponent i resultatpanelen.', 'bot');
      else if (state.step === 'report_done') addMsg('Rapporten \u00e4r klar.', 'bot');
    }
  }
  updatePlaceholder();
}

function createNewProject() {
  toggleProjectMenu();
  try { localStorage.removeItem(_chatStorageKey()); } catch(e) {}
  currentAnalysisId = null;
  chatLog = [];
  state.project = null; state.baseline = null; state.alternatives = null;
  state.selections = {}; state.pendingDesc = null; state.reportMarkdown = null;
  state.chatHistory = []; state.step = 'idle';
  document.getElementById('projectName').textContent = 'Nytt projekt';
  ['projekt','baslinje','alternativ','rapport'].forEach(t => {
    const el = document.getElementById('tab-' + t); if (el) el.disabled = true;
  });
  document.getElementById('resultTabs').style.display = 'none';
  document.getElementById('resultContent').innerHTML = '<div class="empty-state"><p>Beskriv ditt projekt i chatten till vänster för att börja.</p></div>';
  document.getElementById('progressFill').style.width = '0%';
  document.querySelectorAll('.step-circle').forEach((c, i) => { c.className = 'step-circle'; c.textContent = i + 1; });
  document.querySelectorAll('.step-label').forEach(l => l.className = 'step-label');
  const msgs = document.getElementById('messages');
  msgs.innerHTML = '';
  addMsg('Hej! Beskriv ditt projekt. Ange byggnadstyp, byggnadsår, ungefärlig yta och vilka behoven är.', 'bot');
  setLoading(false);
}
</script>
</body>
</html>
"""


def main():
    import argparse
    parser = argparse.ArgumentParser(description='AIda Web UI')
    parser.add_argument('--port', type=int, default=5002)
    parser.add_argument('--host', type=str, default='127.0.0.1')
    args = parser.parse_args()

    print(f"AIda web UI: http://{args.host}:{args.port}")
    app.run(host=args.host, port=args.port, debug=False)


if __name__ == '__main__':
    main()
